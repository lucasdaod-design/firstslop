import streamlit as st
import folium
from streamlit_folium import st_folium
from folium.plugins import MeasureControl
import pandas as pd
import requests
import io
import zipfile
import math
import re
import json
import os
import unicodedata
from urllib.parse import quote
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_OK = True
except Exception:
    DOCX_OK = False

try:
    from geomag.geomag import GeoMag
    GEOMAG_OK = True
except Exception:
    GEOMAG_OK = False


# =====================================================
# CONFIGURAÇÃO
# =====================================================

st.set_page_config(
    page_title="Ferramenta Auxiliar de Planejamento SLOP",
    page_icon="🪂",
    layout="wide"
)

# Mantido para compatibilidade com seu código anterior.
# Nesta versão o mapa usa ESRI World Imagery para evitar cair no OpenStreetMap.
# O aplicativo agora puxa a chave do cofre seguro do Streamlit
GOOGLE_MAPS_API_KEY = st.secrets["GOOGLE_MAPS_API_KEY"]
ARQUIVO_PERFIS_VELAME = "perfis_velame.json"

# =====================================================
# CAIXA PRETA: RADAR DE TRÁFEGO E MISSÕES
# =====================================================
def registrar_log_missao(acao):
    try:
        webhook_url = st.secrets.get("WEBHOOK_GOOGLE_SHEETS")
        if not webhook_url:
            return
        
        # O uso do "f-string" com uma aspa simples (') no início das coordenadas 
        # força o Google Sheets a blindar o número como texto, preservando o ponto decimal!
        payload = {
            "acao": acao,
            "localidade": str(st.session_state.get("localidade_alvo", "-")),
            "lat_alvo": f"'{st.session_state.get('lat', '-')}",
            "lon_alvo": f"'{st.session_state.get('lon', '-')}",
            "lat_aero": f"'{st.session_state.get('lat_aerodromo_partida', '-')}",
            "lon_aero": f"'{st.session_state.get('lon_aerodromo_partida', '-')}"
        }
        
        import threading
        threading.Thread(target=requests.post, args=(webhook_url,), kwargs={"data": payload}).start()
    except Exception:
        pass

# =====================================================
# TABELA NOAA READY — PRESSÃO / ALTITUDE
# =====================================================

# =====================================================
# TABELA NOAA READY — PRESSÃO / ALTITUDE
# =====================================================

PRESSAO_ALTITUDE_NOAA = {
    20: 76500,
    50: 63300,
    100: 51800,
    150: 44300,
    200: 38600,
    250: 34000,
    300: 30000,
    350: 26600,
    400: 23500,
    450: 20800,
    500: 18200,
    550: 15900,
    600: 13800,
    650: 11700,
    700: 9800,
    750: 8000,
    800: 6500,
    850: 4800,
    900: 3250,
    925: 2500,
    950: 1750,
    975: 1000,
    1000: 0,
}

# =====================================================
# ESTADO INICIAL
# =====================================================

if "lat" not in st.session_state:
    st.session_state.lat = -17.4198

if "lon" not in st.session_state:
    st.session_state.lon = -49.5713

if "altitude_ft" not in st.session_state:
    st.session_state.altitude_ft = 1815.0

if "declinacao" not in st.session_state:
    st.session_state.declinacao = None

if "centro_mapa" not in st.session_state:
    st.session_state.centro_mapa = [st.session_state.lat, st.session_state.lon]

if "ultimo_clique_planejamento" not in st.session_state:
    st.session_state.ultimo_clique_planejamento = None

if "pontos_regua" not in st.session_state:
    st.session_state.pontos_regua = []

if "ultimo_clique_regua" not in st.session_state:
    st.session_state.ultimo_clique_regua = None

if "vento_medio_velame" not in st.session_state:
    st.session_state.vento_medio_velame = 0.0

if "direcao_media_velame" not in st.session_state:
    st.session_state.direcao_media_velame = None

if "mapa_planejamento_rev" not in st.session_state:
    st.session_state.mapa_planejamento_rev = 0

if "mapa_regua_rev" not in st.session_state:
    st.session_state.mapa_regua_rev = 0
if "pontos_controle" not in st.session_state:
    st.session_state.pontos_controle = []
if "last_lat_calc" not in st.session_state:
    st.session_state.last_lat_calc = None

if "last_lon_calc" not in st.session_state:
    st.session_state.last_lon_calc = None
if "windgram_texto" not in st.session_state:
    st.session_state.windgram_texto = ""
if "mapa_aerodromo_rev" not in st.session_state:
    st.session_state.mapa_aerodromo_rev = 0

if "ultimo_clique_aerodromo" not in st.session_state:
    st.session_state.ultimo_clique_aerodromo = None

if "lat_aerodromo_partida" not in st.session_state:
    st.session_state.lat_aerodromo_partida = st.session_state.lat

if "lon_aerodromo_partida" not in st.session_state:
    st.session_state.lon_aerodromo_partida = st.session_state.lon

# =====================================================
# FUNÇÕES GERAIS
# =====================================================

def m_para_ft(metros):
    return metros * 3.28084


def nm_para_km(nm):
    return nm * 1.852


def km_para_nm(km):
    return km / 1.852


def m_para_nm(metros):
    return metros / 1852


def kft_para_ft(kft):
    return kft * 1000


def normalizar_azimute(graus):
    return graus % 360


def contra_azimute(graus):
    return (graus + 180) % 360


def verdadeiro_para_magnetico(azimute_verdadeiro, declinacao):
    if declinacao is None:
        declinacao = 0.0
    return normalizar_azimute(azimute_verdadeiro - declinacao)


def media_circular(direcoes):
    if not direcoes:
        return None

    soma_sen = 0
    soma_cos = 0

    for d in direcoes:
        rad = math.radians(d)
        soma_sen += math.sin(rad)
        soma_cos += math.cos(rad)

    media = math.degrees(math.atan2(soma_sen, soma_cos))
    return normalizar_azimute(media)


def media_circular_ponderada(direcoes, velocidades):
    if not direcoes or not velocidades:
        return None

    soma_sen = 0
    soma_cos = 0

    for d, v in zip(direcoes, velocidades):
        rad = math.radians(d)
        soma_sen += math.sin(rad) * v
        soma_cos += math.cos(rad) * v

    media = math.degrees(math.atan2(soma_sen, soma_cos))
    return normalizar_azimute(media)


def normalizar_texto_busca(texto):
    texto = texto.strip().lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = " ".join(texto.split())
    return texto


def buscar_localidade(texto):
    texto_original = texto.strip()

    if not texto_original:
        return None

    texto_norm = normalizar_texto_busca(texto_original)

    apelidos = {
        "campo grande": "Campo Grande, Mato Grosso do Sul, Brasil",
        "campo grande ms": "Campo Grande, Mato Grosso do Sul, Brasil",
        "campo grande, ms": "Campo Grande, Mato Grosso do Sul, Brasil",
        "goiania": "Goiânia, Goiás, Brasil",
        "goiânia": "Goiânia, Goiás, Brasil",
        "goiania go": "Goiânia, Goiás, Brasil",
        "goiânia go": "Goiânia, Goiás, Brasil",
        "anapolis": "Anápolis, Goiás, Brasil",
        "anápolis": "Anápolis, Goiás, Brasil",
        "anapolis go": "Anápolis, Goiás, Brasil",
        "anápolis go": "Anápolis, Goiás, Brasil",
        "rio de janeiro": "Rio de Janeiro, Rio de Janeiro, Brasil",
        "sao paulo": "São Paulo, São Paulo, Brasil",
        "são paulo": "São Paulo, São Paulo, Brasil",
        "brasilia": "Brasília, Distrito Federal, Brasil",
        "brasília": "Brasília, Distrito Federal, Brasil",
    }

    consulta_principal = apelidos.get(texto_norm, texto_original)

    tentativas = [
        consulta_principal,
        f"{texto_original}, Brasil",
        f"Município de {texto_original}, Brasil",
    ]

    # Remove duplicados mantendo a ordem
    tentativas_unicas = []
    for item in tentativas:
        if item not in tentativas_unicas:
            tentativas_unicas.append(item)

    headers = {
        "User-Agent": "CalculadoraPQD/1.0",
        "Accept-Language": "pt-BR,pt;q=0.9,en;q=0.8",
    }

    # =====================================================
    # 1ª tentativa: Nominatim / OpenStreetMap
    # =====================================================
    for consulta in tentativas_unicas:
        try:
            url = "https://nominatim.openstreetmap.org/search"

            params = {
                "q": consulta,
                "format": "json",
                "limit": 5,
                "countrycodes": "br",
                "addressdetails": 1,
            }

            r = requests.get(url, params=params, headers=headers, timeout=15)
            r.raise_for_status()

            dados = r.json()

            if dados:
                escolhido = dados[0]

                return {
                    "nome": escolhido.get("display_name", consulta),
                    "lat": float(escolhido["lat"]),
                    "lon": float(escolhido["lon"]),
                    "fonte": "Nominatim / OpenStreetMap",
                }

        except Exception:
            continue

    # =====================================================
    # 2ª tentativa: Open-Meteo Geocoding
    # =====================================================
    try:
        nome_para_busca = consulta_principal.split(",")[0].strip()

        url = "https://geocoding-api.open-meteo.com/v1/search"

        params = {
            "name": nome_para_busca,
            "count": 10,
            "language": "pt",
            "format": "json",
        }

        r = requests.get(url, params=params, timeout=15)
        r.raise_for_status()

        dados = r.json()
        resultados = dados.get("results", [])

        resultados_br = [
            item for item in resultados
            if item.get("country_code") == "BR"
        ]

        if resultados_br:
            escolhido = resultados_br[0]

            nome = escolhido.get("name", nome_para_busca)
            estado = escolhido.get("admin1", "")
            pais = escolhido.get("country", "Brasil")

            nome_completo = ", ".join(
                parte for parte in [nome, estado, pais] if parte
            )

            return {
                "nome": nome_completo,
                "lat": float(escolhido["latitude"]),
                "lon": float(escolhido["longitude"]),
                "fonte": "Open-Meteo Geocoding",
            }

    except Exception:
        pass

    return None
def calcular_coordenada_destino(lat, lon, distancia_km, azimute_graus):
    R = 6371.0 # Raio médio da Terra em km
    lat1 = math.radians(lat)
    lon1 = math.radians(lon)
    brng = math.radians(azimute_graus)
    
    lat2 = math.asin(math.sin(lat1) * math.cos(distancia_km / R) + 
                     math.cos(lat1) * math.sin(distancia_km / R) * math.cos(brng))
    
    lon2 = lon1 + math.atan2(math.sin(brng) * math.sin(distancia_km / R) * math.cos(lat1), 
                             math.cos(distancia_km / R) - math.sin(lat1) * math.sin(lat2))
    
    return math.degrees(lat2), math.degrees(lon2)
def diferenca_angular(a, b):
    """
    Diferença angular mínima entre dois azimutes.
    Ex.: 350 e 10 = 20 graus.
    """
    return abs((a - b + 180) % 360 - 180)


def detectar_dog_leg(df_windgram, limite_graus=90):
    """
    Procura uma quebra >= limite_graus dentro das camadas de Velame aberto.
    Retorna os dois blocos (superior e inferior) se encontrar.
    """
    if df_windgram is None or df_windgram.empty:
        return None

    base = df_windgram[df_windgram["Fase"] == "Velame aberto"].copy()

    if len(base) < 2:
        return None

    # Ordem do salto: de cima para baixo
    base = base.sort_values("Altitude NOAA ft", ascending=False).reset_index(drop=True)

    for i in range(len(base) - 1):
        dir_atual = float(base.loc[i, "Direção °"])
        dir_prox = float(base.loc[i + 1, "Direção °"])

        diff = diferenca_angular(dir_atual, dir_prox)

        if diff >= limite_graus:
            bloco_superior = base.iloc[:i + 1].copy()
            bloco_inferior = base.iloc[i + 1:].copy()

            dir_superior = media_circular(
                bloco_superior["Direção °"].astype(float).tolist()
            )
            dir_inferior = media_circular(
                bloco_inferior["Direção °"].astype(float).tolist()
            )

            return {
                "houve_dog_leg": True,
                "indice_quebra": i,
                "diferenca_graus": diff,
                "qtd_total": len(base),
                "qtd_superior": len(bloco_superior),
                "qtd_inferior": len(bloco_inferior),
                "dir_superior": dir_superior,
                "dir_inferior": dir_inferior,
                "bloco_superior": bloco_superior,
                "bloco_inferior": bloco_inferior,
            }

    return None


def resolver_geometria_dog_leg(
    lat_alvo,
    lon_alvo,
    dist_total_km,
    azimute_vermelho,
    azimute_inferior,
    azimute_superior,
    qtd_inferior,
    qtd_superior,
):
    """
    Resolve o ponto de quebra do Dog Leg para que:
    - a soma vetorial dos 2 trechos bata com a reta vermelha;
    - o KMZ possa mostrar as 2 retas azuis.
    """

    def vetor_unitario(azimute):
        rad = math.radians(azimute)
        # eixo x = leste ; eixo y = norte
        return math.sin(rad), math.cos(rad)

    ux_inf, uy_inf = vetor_unitario(azimute_inferior)
    ux_sup, uy_sup = vetor_unitario(azimute_superior)

    vx, vy = vetor_unitario(azimute_vermelho)
    vx *= dist_total_km
    vy *= dist_total_km

    det = (ux_inf * uy_sup) - (uy_inf * ux_sup)

    # fallback por proporção de camadas
    def fallback():
        total = qtd_inferior + qtd_superior
        if total <= 0:
            total = 1

        dist_inf = dist_total_km * (qtd_inferior / total)
        dist_sup = dist_total_km * (qtd_superior / total)

        lat_quebra, lon_quebra = calcular_coordenada_destino(
            lat_alvo, lon_alvo, dist_inf, azimute_inferior
        )

        lat_ps_calc, lon_ps_calc = calcular_coordenada_destino(
            lat_quebra, lon_quebra, dist_sup, azimute_superior
        )

        return {
            "dist_inferior_km": dist_inf,
            "dist_superior_km": dist_sup,
            "lat_quebra": lat_quebra,
            "lon_quebra": lon_quebra,
            "lat_ps_calc": lat_ps_calc,
            "lon_ps_calc": lon_ps_calc,
            "metodo": "proporcional",
        }

    if abs(det) < 1e-9:
        return fallback()

    dist_inf = (vx * uy_sup - vy * ux_sup) / det
    dist_sup = (ux_inf * vy - uy_inf * vx) / det

    if dist_inf <= 0 or dist_sup <= 0:
        return fallback()

    lat_quebra, lon_quebra = calcular_coordenada_destino(
        lat_alvo, lon_alvo, dist_inf, azimute_inferior
    )

    lat_ps_calc, lon_ps_calc = calcular_coordenada_destino(
        lat_quebra, lon_quebra, dist_sup, azimute_superior
    )

    return {
        "dist_inferior_km": dist_inf,
        "dist_superior_km": dist_sup,
        "lat_quebra": lat_quebra,
        "lon_quebra": lon_quebra,
        "lat_ps_calc": lat_ps_calc,
        "lon_ps_calc": lon_ps_calc,
        "metodo": "vetorial",
    }

def buscar_altitude(lat, lon):
    try:
        url = f"https://api.open-elevation.com/api/v1/lookup?locations={lat},{lon}"
        r = requests.get(url, timeout=10)
        r.raise_for_status()
        dados = r.json()
        elev_m = float(dados["results"][0]["elevation"])
        return m_para_ft(elev_m)
    except Exception:
        return None
    
def consultar_terreno_e_pressao(lat, lon):
    altitude_ft = None
    qfe_hpa = None

    try:
        url_elev = (
            "https://api.open-meteo.com/v1/elevation"
            f"?latitude={lat}&longitude={lon}"
        )

        resp_elev = requests.get(url_elev, timeout=10)
        resp_elev.raise_for_status()
        dados_elev = resp_elev.json()

        if "elevation" in dados_elev and len(dados_elev["elevation"]) > 0:
            altitude_m = float(dados_elev["elevation"][0])
            altitude_ft = m_para_ft(altitude_m)

    except Exception:
        altitude_ft = None

    try:
        url_pressao = (
            "https://api.open-meteo.com/v1/forecast"
            f"?latitude={lat}&longitude={lon}"
            "&current=surface_pressure"
        )

        resp_pressao = requests.get(url_pressao, timeout=10)
        resp_pressao.raise_for_status()
        dados_pressao = resp_pressao.json()

        if "current" in dados_pressao and "surface_pressure" in dados_pressao["current"]:
            qfe_hpa = float(dados_pressao["current"]["surface_pressure"])

    except Exception:
        qfe_hpa = None

    return altitude_ft, qfe_hpa


def calcular_declinacao(lat, lon, altitude_ft):
    if not GEOMAG_OK:
        return None

    try:
        gm = GeoMag()
        resultado = gm.GeoMag(lat, lon, altitude_ft)
        return float(resultado.dec)
    except Exception:
        return None


def haversine_nm(lat1, lon1, lat2, lon2):
    raio_terra_m = 6371000

    phi1 = math.radians(lat1)
    phi2 = math.radians(lat2)

    dphi = math.radians(lat2 - lat1)
    dlambda = math.radians(lon2 - lon1)

    a = (
        math.sin(dphi / 2) ** 2
        + math.cos(phi1) * math.cos(phi2) * math.sin(dlambda / 2) ** 2
    )

    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    distancia_m = raio_terra_m * c

    return distancia_m / 1852


def distancia_total_nm(pontos):
    if len(pontos) < 2:
        return 0.0

    total = 0.0

    for i in range(len(pontos) - 1):
        total += haversine_nm(
            pontos[i][0],
            pontos[i][1],
            pontos[i + 1][0],
            pontos[i + 1][1],
        )

    return total


# =====================================================
# WINDGRAM NOAA READY
# =====================================================

def detectar_colunas_fhr(texto):
    """
    Detecta linha:
    FHR: + 0. + 3. + 6.
    """

    for linha in texto.splitlines():
        if "FHR" in linha.upper():
            nums = re.findall(r"\+\s*(\d+)", linha)
            if nums:
                return [f"+{n}" for n in nums]

    return []


def processar_windgram(texto):
    """
    Lê linhas:
    700.mb 290@030
    700.mb 290@030 291@031 292@032
    """

    colunas = detectar_colunas_fhr(texto)

    padrao_linha = re.compile(
        r"^\s*(\d{2,4})\.?\s*mb\s+(.+)$",
        re.IGNORECASE
    )

    padrao_vento = re.compile(
        r"(\d{1,3})\s*@\s*(\d{1,3})"
    )

    registros = []

    for linha in texto.splitlines():
        linha = linha.strip()

        m = padrao_linha.match(linha)

        if not m:
            continue

        pressao = int(m.group(1))
        resto = m.group(2)

        if pressao not in PRESSAO_ALTITUDE_NOAA:
            continue

        pares = padrao_vento.findall(resto)

        if not pares:
            continue

        item = {
            "pressao_mb": pressao,
            "altitude_ft_msl": PRESSAO_ALTITUDE_NOAA[pressao],
            "valores": {}
        }

        for idx, par in enumerate(pares):
            direcao = int(par[0])
            velocidade = int(par[1])

            if colunas and idx < len(colunas):
                coluna = colunas[idx]
            else:
                coluna = f"+{idx * 3}"

            item["valores"][coluna] = {
                "direcao_graus": direcao,
                "velocidade_kt": velocidade,
            }

        registros.append(item)

    return registros, colunas


def montar_dataframe_windgram(
    registros,
    coluna,
    altitude_alvo_ft,
    altura_comandamento_ft,
    perda_comandamento_ft,
    altura_saida_ft=None,
):
    linhas = []

    topo_velame_ft_msl = altitude_alvo_ft + max(altura_comandamento_ft - perda_comandamento_ft, 0)
    topo_comandamento_ft_msl = altitude_alvo_ft + altura_comandamento_ft

    if altura_saida_ft is not None and altura_saida_ft > altura_comandamento_ft:
        topo_saida_ft_msl = altitude_alvo_ft + altura_saida_ft
    else:
        topo_saida_ft_msl = topo_comandamento_ft_msl

    # ==========================================================
    # LÓGICA DE PAPIRO: METADE DO INTERVALO (VENTO DE SOLO)
    # ==========================================================
    limite_vento_solo_bruto = altitude_alvo_ft + 750
    altitudes_noaa = sorted([r["altitude_ft_msl"] for r in registros])
    limite_vento_solo_eficaz = limite_vento_solo_bruto
    
    for i in range(len(altitudes_noaa) - 1):
        l_inf = altitudes_noaa[i]
        l_sup = altitudes_noaa[i+1]
        
        # Verifica em qual intervalo o ZL + 750 caiu
        if l_inf <= limite_vento_solo_bruto <= l_sup:
            meio_do_intervalo = (l_inf + l_sup) / 2.0
            if limite_vento_solo_bruto > meio_do_intervalo:
                # Passou da metade: desconsidera a de baixo (Vento de solo vai até l_inf)
                limite_vento_solo_eficaz = l_inf
            else:
                # Não passou da metade: mantém a de baixo no Velame (Vento de solo corta antes)
                limite_vento_solo_eficaz = l_inf - 1
            break

    for r in registros:
        if coluna not in r["valores"]:
            continue

        alt = r["altitude_ft_msl"]
        valor = r["valores"][coluna]

        altura_sobre_alvo = alt - altitude_alvo_ft

        # Classificação das fases com a nova regra
        if alt < altitude_alvo_ft:
            fase = "Abaixo do alvo"
        elif alt <= limite_vento_solo_eficaz:
            fase = "Vento de solo (desprezado)"
        elif limite_vento_solo_eficaz < alt <= topo_velame_ft_msl:
            fase = "Velame aberto"
        elif topo_velame_ft_msl < alt <= topo_comandamento_ft_msl:
            fase = "Comandamento"
        elif topo_comandamento_ft_msl < alt <= topo_saida_ft_msl:
            fase = "Queda livre"
        else:
            fase = "Acima da saída"

        linhas.append({
            "Altitude NOAA ft": alt,
            "Altura sobre alvo ft": altura_sobre_alvo,
            "Pressão mb": r["pressao_mb"],
            "Direção °": valor["direcao_graus"],
            "Velocidade kt": valor["velocidade_kt"],
            "Fase": fase,
        })

    df = pd.DataFrame(linhas)

    if not df.empty:
        df = df.sort_values("Altitude NOAA ft", ascending=False)

    return df


def resumo_por_fase(df, fase):
    if df.empty:
        return None

    base = df[df["Fase"] == fase].copy()

    if base.empty:
        return None

    direcoes = base["Direção °"].astype(float).tolist()
    velocidades = base["Velocidade kt"].astype(float).tolist()

    direcao_aritmetica = sum(direcoes) / len(direcoes)
    direcao_circular = media_circular(direcoes)
    direcao_ponderada = media_circular_ponderada(direcoes, velocidades)

    return {
        "qtd": len(base),

        # Velocidade média dos ventos: média aritmética das velocidades
        "vento_medio": sum(velocidades) / len(velocidades),

        # Direção média simples: média aritmética dos azimutes
        "direcao_media_aritmetica": direcao_aritmetica,

        # Direção correta para azimute/referência: média circular
        "direcao_media_circular": direcao_circular,

        # Mantém compatibilidade com partes antigas do app
        "direcao_media": direcao_circular,

        # Média ponderada pela velocidade
        "direcao_ponderada": direcao_ponderada,
    }


def calcular_media_camadas_mais_baixas(df, qtd_camadas):
    if df.empty:
        return None, pd.DataFrame()

    base = df[df["Fase"] == "Velame aberto"].copy()

    if base.empty:
        return None, pd.DataFrame()

    base = base.sort_values("Altitude NOAA ft", ascending=True)

    qtd_camadas = int(qtd_camadas)

    if qtd_camadas < 1:
        qtd_camadas = 1

    if qtd_camadas > len(base):
        qtd_camadas = len(base)

    selecionadas = base.head(qtd_camadas).copy()

    direcoes = selecionadas["Direção °"].astype(float).tolist()
    velocidades = selecionadas["Velocidade kt"].astype(float).tolist()

    direcao_aritmetica = sum(direcoes) / len(direcoes)
    direcao_circular = media_circular(direcoes)
    direcao_ponderada = media_circular_ponderada(direcoes, velocidades)

    resumo = {
        "qtd_camadas": qtd_camadas,

        # Velocidade média: média aritmética
        "vento_medio": sum(velocidades) / len(velocidades),

        # Direção média simples
        "direcao_media_aritmetica": direcao_aritmetica,

        # Direção circular
        "direcao_media_circular": direcao_circular,

        # Mantém compatibilidade com partes antigas do app
        "direcao_media": direcao_circular,

        # Direção ponderada pela velocidade
        "direcao_ponderada": direcao_ponderada,
    }

    return resumo, selecionadas
def estilo_fases(row):
    fase = row["Fase"]

    if fase == "Vento de solo (desprezado)":
        cor = "background-color: #f4cccc"
    elif fase == "Velame aberto":
        cor = "background-color: #a6ff4d"
    elif fase == "Comandamento":
        cor = "background-color: #ffd966"
    elif fase == "Queda livre":
        cor = "background-color: #6fa8dc"
    elif fase == "Acima da saída":
        cor = "background-color: #dddddd"
    else:
        cor = "background-color: #eeeeee"

    estilo = f"{cor}; color: #000000; font-weight: 600"

    return [estilo for _ in row]

# =====================================================
# FÓRMULAS A E D
# =====================================================

def calcular_d(A_kft, FS_kft, CteHz_kt, V_kt, K_kft_h):
    if A_kft <= FS_kft:
        raise ValueError("A precisa ser maior que FS.")

    if K_kft_h <= 0:
        raise ValueError("K precisa ser maior que zero.")

    vel_efetiva = CteHz_kt + V_kt

    if vel_efetiva <= 0:
        raise ValueError("CteHz + V precisa ser maior que zero.")

    return (A_kft - FS_kft) * vel_efetiva / K_kft_h


def calcular_a(D_nm, FS_kft, CteHz_kt, V_kt, K_kft_h):
    vel_efetiva = CteHz_kt + V_kt

    if vel_efetiva <= 0:
        raise ValueError("CteHz + V precisa ser maior que zero.")

    return FS_kft + (D_nm * K_kft_h) / vel_efetiva
# =====================================================
# PERFIS DE VELAME
# =====================================================

def perfis_padrao_velame():
    return [
        {
            "nome": "Perfil padrão - MMS 350 BT80",
            "constante_horizontal_kt": 23.3,
            "k_kft_h": 40.6,
            "fs_kft": 2.0,
        }
    ]


def carregar_perfis_velame():
    if not os.path.exists(ARQUIVO_PERFIS_VELAME):
        return perfis_padrao_velame()

    try:
        with open(ARQUIVO_PERFIS_VELAME, "r", encoding="utf-8") as arquivo:
            perfis = json.load(arquivo)

        if isinstance(perfis, list) and len(perfis) > 0:
            return perfis

        return perfis_padrao_velame()

    except Exception:
        return perfis_padrao_velame()


def salvar_perfis_velame(perfis):
    with open(ARQUIVO_PERFIS_VELAME, "w", encoding="utf-8") as arquivo:
        json.dump(perfis, arquivo, ensure_ascii=False, indent=4)
def formatar_coord_dm(lat, lon):
    def converter(valor, tipo):
        hemisferio = ""

        if tipo == "lat":
            hemisferio = "S" if valor < 0 else "N"
        else:
            hemisferio = "O" if valor < 0 else "L"

        valor_abs = abs(valor)
        graus = int(valor_abs)
        minutos = (valor_abs - graus) * 60

        return f"{graus}° {minutos:.3f}'{hemisferio}"

    return f"{converter(lat, 'lat')}   {converter(lon, 'lon')}"
def formatar_lat_dm(lat):
    valor_abs = abs(lat)
    graus = int(valor_abs)
    minutos = (valor_abs - graus) * 60
    hemisferio = "S" if lat < 0 else "N"
    return f"{graus}° {minutos:.3f}'{hemisferio}"


def formatar_lon_dm(lon):
    valor_abs = abs(lon)
    graus = int(valor_abs)
    minutos = (valor_abs - graus) * 60
    hemisferio = "O" if lon < 0 else "L"
    return f"{graus}° {minutos:.3f}'{hemisferio}"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def gerar_folder_piloto_docx(dados):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("FOLDER DO PILOTO")
    run.bold = True
    run.font.size = Pt(14)

    # Espaço reservado para o usuário colar o mapa manualmente
    doc.add_paragraph("")
    p_aviso = doc.add_paragraph()
    p_aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_aviso = p_aviso.add_run("[ Cole o print do Google Earth aqui ]")
    run_aviso.font.italic = True
    doc.add_paragraph("")

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    linhas = [
        (f"ZL / Local: {dados.get('localidade', '-')}\nCoord: {dados['coord_zl']}", f"Altitude da ZL: {dados['altitude_zl_ft']}"),
        (f"Eixo de lançamento: {dados['eixo_lancamento']}", f"Altitude Adrm: {dados['altitude_aerodromo_ft']}"),
        (f"Eixo de navegação do saltador: {dados['eixo_navegacao']}", f"Altitude PS: {dados['altitude_ps_ft']} / DAA: {dados['daa_qfe']}"),
        (f"Alt comandamento: {dados['altura_comandamento_ft']}", f"Ajuste de altímetro: {dados['ajuste_altimetro']}"),
        (f"Velocidade da Anv: {dados['velocidade_anv']}", "Pqdt embarcados:")
    ]

    for i, linha in enumerate(linhas):
        for j, texto in enumerate(linha):
            cell = table.cell(i, j)
            set_cell_text(cell, texto, bold=False, size=12)
            if i in [0, 2]:
                set_cell_shading(cell, "D9D9D9")
        doc.add_paragraph("")

    tabela_ref = doc.add_table(rows=4, cols=3)
    tabela_ref.style = "Table Grid"
    tabela_ref.alignment = WD_TABLE_ALIGNMENT.CENTER

    cores_colunas = ["E6332A", "FFFF66", "63B35D"]
    cabecalhos = ["4' FORA", "1' FORA", "PONTO DE SAÍDA - PS"]
    
    linhas_ref = [
        cabecalhos,
        [f"Latitude: {dados.get('lat_4_fora', '-')}", f"Latitude: {dados.get('lat_1_fora', '')}", f"Latitude: {dados.get('ps_lat_dm', '')}"],
        [f"Longitude: {dados.get('lon_4_fora', '-')}", f"Longitude: {dados.get('lon_1_fora', '')}", f"Longitude: {dados.get('ps_lon_dm', '')}"],
        ["LUZ VERMELHA", "REPORTAR ANV NA FINAL", "LUZ VERDE APÓS O N/A ROTA OU SOBRE O PS"]
    ]

    for i, linha in enumerate(linhas_ref):
        for j, texto in enumerate(linha):
            cell = tabela_ref.cell(i, j)
            set_cell_text(cell, texto, bold=(i == 0 or i == 3), size=12)
            set_cell_shading(cell, cores_colunas[j])
            for paragrafo in cell.paragraphs:
                if i == 0 or i == 3:
                    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
# =====================================================
# MAPA AUXILIAR
# =====================================================

def criar_mapa_base(location, zoom=13):
    mapa = folium.Map(
        location=location,
        zoom_start=zoom,
        tiles=None,
        control_scale=True,
    )

    # Camada base: satélite
    folium.TileLayer(
        tiles="https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}",
        attr="Esri World Imagery",
        name="Satélite",
        overlay=False,
        control=False,
        show=True,
    ).add_to(mapa)

    # Camada sobreposta: nomes de cidades, limites e referências
    folium.TileLayer(
        tiles="https://services.arcgisonline.com/ArcGIS/rest/services/Reference/World_Boundaries_and_Places/MapServer/tile/{z}/{y}/{x}",
        attr="Esri Reference",
        name="Nomes e limites",
        overlay=True,
        control=False,
        show=True,
    ).add_to(mapa)

    return mapa

# =====================================================
# INTERFACE E TÍTULO
# =====================================================

# Dei um pouco mais de espaço para o texto (3.2) para não sobrepor a caveira
col_titulo1, col_titulo2, col_titulo3 = st.columns([0.5, 3.2, 1.1], vertical_alignment="center")

with col_titulo1:
    st.markdown("<h1 style='text-align: center; font-size: 3rem; margin: 0;'>🪂</h1>", unsafe_allow_html=True)

with col_titulo2:
    # A mágica aqui é o 'white-space: nowrap;', que proíbe a quebra de linha!
    st.markdown("<h1 style='text-align: left; margin: 0; white-space: nowrap;'>Ferramenta Auxiliar de Planejamento SLOP</h1>", unsafe_allow_html=True)

with col_titulo3:
    try:
        st.image("firstsemfundocerto.png", width=85) 
    except Exception:
        st.warning("⚠️ Logo não encontrado")

st.caption(
    "Fluxo: selecionar ponto → abrir NOAA READY → colar Windgram → calcular dados auxiliares."
)

aba_planejamento, aba_calculos, aba_camadas, aba_dkva, aba_folder = st.tabs(
    [
        "Planejamento / Windgram",
        "Cálculo da Distância para Velame Aberto",
        "Calculadora dos Pontos de Controle",
        "Salto sobre o Alvo (DKVA)",
        "Folder do Piloto"
    ]
)


# =====================================================
# ABA PLANEJAMENTO
# =====================================================

with aba_planejamento:
    col_esq, col_dir = st.columns([1.15, 1])

    with col_esq:
        st.subheader("1. Alvo")

        # --- NOVO: ZLs PRÉ-CADASTRADAS ---
        zls_cadastradas = {
            "Personalizado / Outro": None,
            "Itograss": {"lat": -16.479218, "lon": -49.226422},
            "Skydive Cerrado": {"lat": -16.362042, "lon": -48.928371}
        }
        
        zl_selecionada = st.selectbox(
            "Zonas de Lançamento Rápidas", 
            list(zls_cadastradas.keys()),
            help="Selecione uma ZL salva para preencher as coordenadas e o mapa automaticamente."
        )
        
        if zl_selecionada != "Personalizado / Outro" and st.session_state.get("ultima_zl_selecionada") != zl_selecionada:
            st.session_state.lat = zls_cadastradas[zl_selecionada]["lat"]
            st.session_state.lon = zls_cadastradas[zl_selecionada]["lon"]
            st.session_state.localidade_alvo = zl_selecionada
            st.session_state.centro_mapa = [st.session_state.lat, st.session_state.lon]
            st.session_state.mapa_planejamento_rev += 1
            st.session_state.ultima_zl_selecionada = zl_selecionada
            # ❌ st.rerun() FOI EXTIRPADO DAQUI
            
        elif zl_selecionada == "Personalizado / Outro":
            st.session_state.ultima_zl_selecionada = "Personalizado / Outro"

        # Caixa para digitar o nome da ZL (AGORA COM CHAVE DINÂMICA)
        localidade_alvo = st.text_input(
            "Localidade / Nome da ZL", 
            value=st.session_state.get("localidade_alvo", ""),
            placeholder="Ex: ZL Boi Preto / Anápolis-GO",
            key=f"input_local_alvo_{st.session_state.mapa_planejamento_rev}"
        )
        st.session_state.localidade_alvo = localidade_alvo

        busca = st.text_input(
            "Pesquisar coordenadas por cidade",
            placeholder="Ex: Campo Grande, Goiânia, Anápolis, Fazenda..."
        )

        if st.button("🔎 Pesquisar local"):
            if busca.strip():
                resultado = buscar_localidade(busca)
                if resultado:
                    st.session_state.lat = float(resultado["lat"])
                    st.session_state.lon = float(resultado["lon"])
                    st.session_state.centro_mapa = [st.session_state.lat, st.session_state.lon]
                    st.session_state.mapa_planejamento_rev += 1
                    st.session_state.localidade_alvo = resultado["nome"]
                    st.success(resultado["nome"])
                    # ❌ st.rerun() FOI EXTIRPADO DAQUI TAMBÉM
                else:
                    st.error("Local não encontrado.")

        # =====================================================
        # GATILHO AUTOMÁTICO DE ALTITUDE E DECLINAÇÃO
        # =====================================================
        if (st.session_state.get("last_lat_calc") != st.session_state.lat or 
            st.session_state.get("last_lon_calc") != st.session_state.lon):
            
            with st.spinner("📍 Calculando altitude e declinação automática..."):
                alt = buscar_altitude(st.session_state.lat, st.session_state.lon)
                if alt is not None:
                    st.session_state.altitude_ft = round(alt, 0)

                dec = calcular_declinacao(st.session_state.lat, st.session_state.lon, st.session_state.altitude_ft)
                if dec is not None:
                    st.session_state.declinacao = dec
                    
            # Grava na memória para não rodar em loop
            st.session_state.last_lat_calc = st.session_state.lat
            st.session_state.last_lon_calc = st.session_state.lon
            # ❌ O ÚLTIMO st.rerun() FOI EXTIRPADO DAQUI!
        c1, c2, c3 = st.columns(3)

        with c1:
            lat = st.number_input(
                "Latitude",
                value=float(st.session_state.lat),
                step=0.0001,
                format="%.6f",
                key=f"lat_planejamento_{st.session_state.mapa_planejamento_rev}",
            )
            st.session_state.lat = float(lat)
            
        with c2:
            lon = st.number_input(
                "Longitude",
                value=float(st.session_state.lon),
                step=0.0001,
                format="%.6f",
                key=f"lon_planejamento_{st.session_state.mapa_planejamento_rev}",
            )
            st.session_state.lon = float(lon)
            
        with c3:
            altitude_alvo_ft = st.number_input(
                "Altitude do alvo ft MSL",
                value=float(st.session_state.altitude_ft),
                step=50.0,
                key=f"altitude_planejamento_{st.session_state.mapa_planejamento_rev}",
            )
            st.session_state.altitude_ft = float(altitude_alvo_ft)

        st.session_state.centro_mapa = [st.session_state.lat, st.session_state.lon]

        if st.session_state.declinacao is None:
            declinacao = st.number_input(
                "Declinação magnética manual",
                value=0.0,
                step=0.1,
                key="declinacao_manual",
            )
        else:
            declinacao = st.session_state.declinacao

        st.metric("Declinação magnética", f"{declinacao:.2f}°")

        st.markdown("### Selecionar ponto no mapa")

        mapa_planejamento = criar_mapa_base(
            [st.session_state.lat, st.session_state.lon], zoom=13
        )

        folium.Marker(
            location=[st.session_state.lat, st.session_state.lon],
            popup="Alvo",
            tooltip="Alvo",
            icon=folium.Icon(color="red", icon="flag"),
        ).add_to(mapa_planejamento)

        resultado_mapa_planejamento = st_folium(
            mapa_planejamento,
            width=None,
            height=420,
            key=f"mapa_planejamento_{st.session_state.mapa_planejamento_rev}",
            returned_objects=["last_clicked"],
        )

        if resultado_mapa_planejamento and resultado_mapa_planejamento.get("last_clicked"):
            lat_click = float(resultado_mapa_planejamento["last_clicked"]["lat"])
            lon_click = float(resultado_mapa_planejamento["last_clicked"]["lng"])

            novo_clique = [round(lat_click, 7), round(lon_click, 7)]

            if st.session_state.ultimo_clique_planejamento != novo_clique:
                st.session_state.ultimo_clique_planejamento = novo_clique
                st.session_state.lat = lat_click
                st.session_state.lon = lon_click
                st.session_state.centro_mapa = [lat_click, lon_click]
                st.session_state.mapa_planejamento_rev += 1
                st.rerun()

    st.divider()
    st.subheader("2. Parâmetros")

    c4, c5, c6 = st.columns(3)

    with c4:
            altura_comandamento_ft = st.number_input(
                "Altura de comandamento ft",
                value=12000.0,
                step=500.0,
                key="altura_comandamento",
            )
    st.session_state.altura_comandamento_ft = altura_comandamento_ft

    with c5:
            perda_comandamento_ft = st.number_input(
                "Perda no comandamento ft",
                value=1000.0,
                step=100.0,
                key="perda_comandamento",
            )

    with c6:
            altura_saida_ft = st.number_input(
                "Altura de saída ft",
                value=12000.0,
                step=500.0,
                help="Opcional. Se for igual ao comandamento, praticamente não haverá faixa de queda livre colorida.",
                key="altura_saida",
            )

    topo_velame_msl = altitude_alvo_ft + max(altura_comandamento_ft - perda_comandamento_ft, 0)
    topo_comandamento_msl = altitude_alvo_ft + altura_comandamento_ft

        
        # =====================================================
        # PAINEL DINÂMICO DO PONTO DE SAÍDA (PS)
        # =====================================================
    
    st.divider()

    st.subheader("3. Windgram / Dados de Vento")

    st.info("💡 Aqui você deve copiar e Colar o Windgrama com as três colunas que aparecem no NOAA, do jeito que elas aparecem! É só copiar e colar aqui!")

    st.link_button(
                "Abrir NOAA READY",
                "https://www.ready.noaa.gov/READYcmet.php"
            )

    st.write("Coordenadas para copiar no NOAA (copie uma por vez):")
    col_copy1, col_copy2 = st.columns(2)
    
    with col_copy1:
        st.caption("Latitude:")
        st.code(f"{st.session_state.lat:.6f}")
        
    with col_copy2:
        st.caption("Longitude:")
        st.code(f"{st.session_state.lon:.6f}")
    with st.expander("Instruções rápidas"):
                st.markdown(
                    """
                    No NOAA READY:

                    1. Insira latitude e longitude.
                    2. Escolha a previsão mais recente.
                    3. Produto: **WINDGRAM**.
                    4. Escolha o horário Zulu do lançamento.
                    5. Duração: **3h**.
                    6. Saída: **Text only**.
                    7. Resolva o CAPTCHA no site.
                    8. Copie a tabela textual e cole abaixo.
                    """
                )

    texto_windgram = st.text_area(
                "Cole aqui o Windgram textual",
                height=260,
                placeholder="Ex:\nFHR: + 0.\n700.mb 290@030\n750.mb 289@029\n800.mb 286@026",
                key="windgram_texto",
            )

    if st.button("Carregar Windgram e calcular", type="primary"):
                registros, colunas = processar_windgram(texto_windgram)

                if not registros:
                    st.error("Não consegui ler o Windgram. Cole linhas no formato: 700.mb 290@030")
                else:
                    if not colunas:
                        colunas = list(registros[0]["valores"].keys())

                    coluna = colunas[0]

                    df = montar_dataframe_windgram(
                        registros=registros,
                        coluna=coluna,
                        altitude_alvo_ft=altitude_alvo_ft,
                        altura_comandamento_ft=altura_comandamento_ft,
                        perda_comandamento_ft=perda_comandamento_ft,
                        altura_saida_ft=altura_saida_ft,
                    )

                    resumo_velame = resumo_por_fase(df, "Velame aberto")
                    resumo_queda = resumo_por_fase(df, "Queda livre")

                    st.session_state.df_windgram = df
                    st.session_state.resumo_velame = resumo_velame
                    st.session_state.resumo_queda = resumo_queda
                    st.session_state.declinacao_usada = declinacao

                    if resumo_velame:
                        st.session_state.vento_medio_velame = resumo_velame["vento_medio"]

                        st.session_state.direcao_vento_verdadeira_kmz = float(
                            resumo_velame.get(
                                "direcao_ponderada",
                                resumo_velame.get("direcao_media", 0.0)
                            )
                        )

                        st.session_state.direcao_media_velame = float(
                            resumo_velame.get(
                                "direcao_media_circular",
                                resumo_velame.get("direcao_media", 0.0)
                            )
                        )
                    st.success("Windgram processado com sucesso.")

    with col_dir:
            st.subheader("Resultado")

            if "df_windgram" not in st.session_state:
                st.info("Cole o Windgram e clique em calcular.")
            else:
                df = st.session_state.df_windgram
                resumo_velame = st.session_state.resumo_velame
                resumo_queda = st.session_state.resumo_queda
                declinacao = st.session_state.declinacao_usada

                if resumo_velame:
                    vento = resumo_velame["vento_medio"]

                    direcao_vento_aritmetica = resumo_velame.get(
                        "direcao_media_aritmetica",
                        resumo_velame.get("direcao_media", 0.0)
                    )

                    direcao_vento_circular = resumo_velame.get(
                        "direcao_media_circular",
                        resumo_velame.get("direcao_media", 0.0)
                    )

                    direcao_ponderada = float(
                        resumo_velame.get(
                            "direcao_ponderada",
                            resumo_velame.get("direcao_media", 0.0)
                        )
                    )

                    # Base verdadeira/geográfica para os cálculos: Direção Ponderada
                    azimute_referencia_verdadeiro = direcao_ponderada

                    azimute_referencia_magnetico = verdadeiro_para_magnetico(
                        azimute_referencia_verdadeiro,
                        declinacao
                    )

                    # --- ALERTA DE VENTO FORTE (> 18 kt) ---
                    camadas_fortes = df[(df["Fase"] == "Velame aberto") & (df["Velocidade kt"] > 18)]
                    if not camadas_fortes.empty:
                        vento_max = camadas_fortes["Velocidade kt"].max()
                        st.error(f"⚠️ **ATENÇÃO OPERACIONAL:** Vento forte detectado na camada de velame aberto (Pico de {vento_max:.0f} kt). Avalie as margens de segurança para o salto!")
                    # ---------------------------------------

                    r1, r2 = st.columns(2)

                    with r1:
                        st.metric(
                            "Média dos Ventos de Camada",
                            f"{vento:.1f} kt",
                            help="Média aritmética das velocidades dos ventos."
                        )

                        st.metric(
                            "Direção Ponderada dos Ventos",
                            f"{direcao_ponderada:.0f}°",
                            help="Média vetorial ponderada pela força do vento (Referência Principal)."
                        )

                    with r2:
                        st.metric(
                            "Entrada de Nariz (Ref: Média PONDERADA)",
                            f"{azimute_referencia_magnetico:.0f}°"
                        )

                        st.metric(
                            "Entrada de Cauda (Ref: Média PONDERADA)",
                            f"{contra_azimute(azimute_referencia_magnetico):.0f}°"
                        )
                        
                    st.write(
                            f"Referência verdadeira — Média Circular (Sem peso): "
                            f"**{direcao_vento_circular:.0f}°**"
                    )
                    st.write(f"Declinação aplicada: **{declinacao:.2f}°**")
                    st.write(f"Direção média aritmética dos ventos: **{direcao_vento_aritmetica:.0f}°**")
                    st.write(f"Camadas de velame usadas: **{resumo_velame['qtd']}**")

                else:
                    st.warning("Nenhuma camada de velame aberto foi encontrada.")

                if resumo_queda:
                    with st.expander("Resumo queda livre"):
                        st.write(f"Vento médio: **{resumo_queda['vento_medio']:.1f} kt**")
                        st.write(f"Direção média: **{resumo_queda['direcao_media']:.0f}°**")
                        st.write(f"Camadas usadas: **{resumo_queda['qtd']}**")

                with st.expander("Ver tabela colorida"):
                    st.dataframe(
                        df.style.apply(estilo_fases, axis=1),
                        use_container_width=True,
                        height=500
                    )


# =====================================================
# ABA CÁLCULO DA DISTÂNCIA PARA VELAME ABERTO
# =====================================================

with aba_calculos:
    st.subheader("Cálculo da Distância para Velame Aberto")

    st.info(
        "Objetivo: calcular D pela fórmula: "
        "D = (A - FS) × (Constante Horizontal + Vento Médio) / K"
    )

    perfis = carregar_perfis_velame()

    nomes_perfis = [p["nome"] for p in perfis]

    perfil_escolhido_nome = st.selectbox(
        "Selecionar perfil salvo",
        nomes_perfis
    )

    perfil_escolhido = next(
        p for p in perfis if p["nome"] == perfil_escolhido_nome
    )

    st.divider()

    vento_default = float(st.session_state.vento_medio_velame)

    if vento_default <= 0:
        st.warning(
            "Ainda não há vento médio vindo da aba Planejamento / Windgram. "
            "Cole e processe o Windgram na aba anterior para preencher automaticamente."
        )

    c1, c2 = st.columns(2)

    with c1:
        altura_comandamento_kft_auto = float(
        st.session_state.get("altura_comandamento_ft", 12000.0)
    ) / 1000.0

        A_kft = st.number_input(
        "A — Altura de comandamento do paraquedas (kft)",
        value=altura_comandamento_kft_auto,
        step=0.1,
        key=f"calc_d_a_auto_{int(st.session_state.get('altura_comandamento_ft', 12000.0))}",
        help="Valor preenchido automaticamente a partir da Altura de comandamento informada na aba Planejamento."
    )
        FS_kft = st.number_input(
            "FS — Fator de Segurança (kft)",
            value=float(perfil_escolhido.get("fs_kft", 2.0)),
            step=0.1,
            help="Exemplo: 2.0 significa 2.000 ft."
        )

        st.warning(
            "Alerta FS: o mínimo do manual é 2 kft — "
            "1 kft para abertura do paraquedas e 1 kft para o Líder chegar com 1000 ft sobre o alvo."
        )

    with c2:
        CteHz = st.number_input(
            "Constante Horizontal (kt)",
            value=float(perfil_escolhido.get("constante_horizontal_kt", 23.3)),
            step=0.1
        )

        V = st.number_input(
            "Velocidade média dos ventos de camada (kt)",
            value=vento_default,
            step=0.1,
            help="Este valor vem automaticamente da aba Planejamento / Windgram."
        )

        K = st.number_input(
            "K — Constante vertical (kft/h)",
            value=float(perfil_escolhido.get("k_kft_h", 40.6)),
            step=0.1
        )

    st.divider()

    if st.button("Calcular D", type="primary"):
        try:
            if FS_kft < 2.0:
                st.error("FS abaixo de 2 kft. Ajuste o fator de segurança antes de calcular.")
            else:
                D = calcular_d(A_kft, FS_kft, CteHz, V, K)

                st.session_state.distancia_velame_aberto_nm = D
                st.session_state.fs_velame_aberto_kft = FS_kft
                st.session_state.ctehz_velame_aberto_kt = CteHz
                st.session_state.k_velame_aberto_kft_h = K
                
                # --- NOVO: SALVAR O PONTO D NA MEMÓRIA IMEDIATAMENTE ---
                azimute_vento = float(st.session_state.get("direcao_vento_verdadeira_kmz", 0.0))
                dist_km = nm_para_km(D)
                lat_d, lon_d = calcular_coordenada_destino(st.session_state.lat, st.session_state.lon, dist_km, azimute_vento)
                
                st.session_state.ps_lat = lat_d
                st.session_state.ps_lon = lon_d
                st.session_state.ps_origem = "Ponto Limite D (Velame Aberto)"
                # ---------------------------------------------------------

                r1, r2, r3 = st.columns(3)

                with r1:
                    st.metric("D em NM", f"{D:.3f} NM")

                with r2:
                    st.metric("D em km", f"{nm_para_km(D):.3f} km")

                with r3:
                    st.metric("D em metros", f"{D * 1852:.0f} m")

                st.success("Distância para velame aberto calculada com sucesso.")

        except Exception as e:
            st.error(str(e))

    st.divider()
    st.markdown("#### 📍 Consulta do PS e Aeródromo")

    # -----------------------------
    # PS
    # -----------------------------

    ps_disponivel = "ps_lat" in st.session_state and "ps_lon" in st.session_state

    if not ps_disponivel:
        st.warning(
            "Coordenada do PS ainda não registrada. "
            "Gere o KMZ ou registre o PS antes de consultar altitude e DAA/QFE do PS."
        )
    else:
        lat_padrao_consulta = float(st.session_state.ps_lat)
        lon_padrao_consulta = float(st.session_state.ps_lon)

        st.markdown("##### Ponto de Saída")

        c_ps_lat, c_ps_lon = st.columns(2)

        with c_ps_lat:
            lat_consulta = st.number_input(
                "Latitude do PS",
                value=lat_padrao_consulta,
                step=0.0001,
                format="%.6f",
                key=f"lat_consulta_ps_{round(lat_padrao_consulta, 6)}"
            )

        with c_ps_lon:
            lon_consulta = st.number_input(
                "Longitude do PS",
                value=lon_padrao_consulta,
                step=0.0001,
                format="%.6f",
                key=f"lon_consulta_ps_{round(lon_padrao_consulta, 6)}"
            )

        st.success("Coordenada do PS carregada automaticamente.")
        st.caption(
            f"PS: {lat_consulta:.6f}, {lon_consulta:.6f} | "
            f"Origem: {st.session_state.get('ps_origem', 'não informada')}"
        )

       # -----------------------------
    # AERÓDROMO
    # -----------------------------

    st.markdown("##### Aeródromo de partida")

    with st.form("form_busca_aerodromo"):
        busca_aerodromo = st.text_input(
            "Buscar aeródromo/localidade de partida",
            placeholder="Ex: Campo Grande MS, Anápolis, Goiânia, Base Aérea..."
        )

        buscar_aerodromo = st.form_submit_button("🔎 Buscar local do aeródromo")

    if buscar_aerodromo:
        if busca_aerodromo.strip():
            resultado_aero = buscar_localidade(busca_aerodromo)

            if resultado_aero:
                st.session_state.lat_aerodromo_partida = float(resultado_aero["lat"])
                st.session_state.lon_aerodromo_partida = float(resultado_aero["lon"])
                st.session_state.mapa_aerodromo_rev += 1

                st.success(
                    f"{resultado_aero['nome']} | Fonte: {resultado_aero.get('fonte', 'Busca online')}"
                )

                st.rerun()

            else:
                st.error(
                    "Aeródromo/localidade não encontrado. Tente escrever com o estado, exemplo: Campo Grande MS."
                )

    c_aero_lat, c_aero_lon = st.columns(2)

    with c_aero_lat:
        lat_aerodromo = st.number_input(
            "Latitude do aeródromo de partida",
            value=float(st.session_state.lat_aerodromo_partida),
            step=0.0001,
            format="%.6f",
            key=f"lat_aerodromo_partida_input_{st.session_state.mapa_aerodromo_rev}"
        )

    with c_aero_lon:
        lon_aerodromo = st.number_input(
            "Longitude do aeródromo de partida",
            value=float(st.session_state.lon_aerodromo_partida),
            step=0.0001,
            format="%.6f",
            key=f"lon_aerodromo_partida_input_{st.session_state.mapa_aerodromo_rev}"
        )

    st.session_state.lat_aerodromo_partida = float(lat_aerodromo)
    st.session_state.lon_aerodromo_partida = float(lon_aerodromo)

    st.markdown("###### Selecionar aeródromo no mapa")

    mapa_aerodromo = criar_mapa_base(
        [
            st.session_state.lat_aerodromo_partida,
            st.session_state.lon_aerodromo_partida
        ],
        zoom=12
    )

    folium.Marker(
        location=[
            st.session_state.lat_aerodromo_partida,
            st.session_state.lon_aerodromo_partida
        ],
        popup="Aeródromo de partida",
        tooltip="Aeródromo de partida",
        icon=folium.Icon(color="blue", icon="flag"),
    ).add_to(mapa_aerodromo)

    resultado_mapa_aerodromo = st_folium(
        mapa_aerodromo,
        width=None,
        height=360,
        key=f"mapa_aerodromo_{st.session_state.mapa_aerodromo_rev}",
        returned_objects=["last_clicked"],
    )

    if resultado_mapa_aerodromo and resultado_mapa_aerodromo.get("last_clicked"):
        lat_click_aero = float(resultado_mapa_aerodromo["last_clicked"]["lat"])
        lon_click_aero = float(resultado_mapa_aerodromo["last_clicked"]["lng"])

        novo_clique_aero = [
            round(lat_click_aero, 7),
            round(lon_click_aero, 7)
        ]

        if st.session_state.ultimo_clique_aerodromo != novo_clique_aero:
            st.session_state.ultimo_clique_aerodromo = novo_clique_aero
            st.session_state.lat_aerodromo_partida = lat_click_aero
            st.session_state.lon_aerodromo_partida = lon_click_aero
            st.session_state.mapa_aerodromo_rev += 1
            st.rerun()

    st.caption(
        f"Aeródromo selecionado: "
        f"{st.session_state.lat_aerodromo_partida:.6f}, "
        f"{st.session_state.lon_aerodromo_partida:.6f}"
    )
    # -----------------------------
    # BOTÕES
    # -----------------------------

    b1, b2 = st.columns(2)

    with b1:
        consultar_ps = st.button(
            "🌎 Consultar altitude e DAA/QFE do PS",
            key="btn_consultar_ambiente_ps"
        )

    with b2:
        calcular_altimetria = st.button(
            "🧮 Calcular Diferença Altimétrica",
            key="btn_calcular_diferenca_altimetrica"
        )

    if consultar_ps:
        if not ps_disponivel:
            st.error("PS ainda não registrado. Gere o KMZ primeiro ou registre o PS.")
        else:
            altitude_ps_ft, qfe_ps_hpa = consultar_terreno_e_pressao(
                lat_consulta,
                lon_consulta
            )

            st.session_state.altitude_consulta_ft = altitude_ps_ft
            st.session_state.qfe_consulta_hpa = qfe_ps_hpa

    if calcular_altimetria:
        altitude_aerodromo_ft, _ = consultar_terreno_e_pressao(
            lat_aerodromo,
            lon_aerodromo
        )

        altitude_alvo_ft_ref = float(st.session_state.get("altitude_ft", 0.0))

        st.session_state.altitude_aerodromo_partida_ft = altitude_aerodromo_ft

        if altitude_aerodromo_ft is not None:
            st.session_state.altimetro_aerodromo_alvo_ft = (
                altitude_aerodromo_ft - altitude_alvo_ft_ref
            )
        else:
            st.session_state.altimetro_aerodromo_alvo_ft = None

    # -----------------------------
    # RESULTADOS
    # -----------------------------

    r1, r2 = st.columns(2)

    with r1:
        altitude_consulta_ft = st.session_state.get("altitude_consulta_ft")

        if altitude_consulta_ft is not None:
            st.metric(
                "Altitude do terreno no PS",
                f"{altitude_consulta_ft:,.0f} ft".replace(",", "X").replace(".", ",").replace("X", ".")
            )
        else:
            st.metric("Altitude do terreno no PS", "—")

    with r2:
        qfe_consulta_hpa = st.session_state.get("qfe_consulta_hpa")

        if qfe_consulta_hpa is not None:
            st.metric("DAA / QFE no PS", f"{qfe_consulta_hpa:.1f} hPa")
        else:
            st.metric("DAA / QFE no PS", "—")

    r3, r4 = st.columns(2)

    with r3:
        altitude_aerodromo_ft = st.session_state.get("altitude_aerodromo_partida_ft")

        if altitude_aerodromo_ft is not None:
            st.metric(
                "Altitude do aeródromo de partida",
                f"{altitude_aerodromo_ft:,.0f} ft".replace(",", "X").replace(".", ",").replace("X", ".")
            )
        else:
            st.metric("Altitude do aeródromo de partida", "—")

    with r4:
        altimetro_ft = st.session_state.get("altimetro_aerodromo_alvo_ft")

        if altimetro_ft is not None:
            st.metric(
                "Diferença Altimétrica",
                f"{altimetro_ft:,.0f} ft".replace(",", "X").replace(".", ",").replace("X", ".")
            )
            st.caption("Altitude do aeródromo de partida - altitude do alvo")
        else:
            st.metric("Diferença Altimétrica", "—")

    st.caption("Fonte de consulta: Open-Meteo Elevation API e Open-Meteo Forecast Surface Pressure.")    
    with st.expander("Salvar novo perfil"):
        with st.form("form_salvar_perfil_velame", clear_on_submit=True):
            novo_nome = st.text_input(
                "Nome do perfil",
                placeholder="Ex: MMS 350, Phantom 360, Perfil teste..."
            )

            novo_ctehz = st.number_input(
                "Constante Horizontal do perfil (kt)",
                value=23.3,
                step=0.1,
                key="novo_perfil_ctehz"
            )

            novo_k = st.number_input(
                "K do perfil (kft/h)",
                value=40.6,
                step=0.1,
                key="novo_perfil_k"
            )

            novo_fs = st.number_input(
                "FS padrão do perfil (kft)",
                value=2.0,
                step=0.1,
                key="novo_perfil_fs"
            )

            salvar = st.form_submit_button("Salvar perfil")

            if salvar:
                if not novo_nome.strip():
                    st.error("Informe um nome para o perfil.")
                else:
                    perfis.append(
                        {
                            "nome": novo_nome.strip(),
                            "constante_horizontal_kt": float(novo_ctehz),
                            "k_kft_h": float(novo_k),
                            "fs_kft": float(novo_fs),
                        }
                    )

                    salvar_perfis_velame(perfis)

                    st.success("Perfil salvo com sucesso.")
                    st.rerun()

    with st.expander("Excluir perfil salvo"):
        if len(perfis) <= 1:
            st.info("Há apenas o perfil padrão. Crie outro perfil antes de excluir.")
        else:
            perfil_excluir = st.selectbox(
                "Escolha o perfil para excluir",
                nomes_perfis,
                key="perfil_excluir"
            )

            if st.button("Excluir perfil"):
                perfis = [p for p in perfis if p["nome"] != perfil_excluir]
                salvar_perfis_velame(perfis)
                st.success("Perfil excluído.")
                st.rerun()

# =====================================================
# ABA CALCULADORA DOS PONTOS DE CONTROLE
# =====================================================

# =====================================================
# ABA CALCULADORA DOS PONTOS DE CONTROLE
# =====================================================

with aba_camadas:
    st.subheader("Calculadora dos Pontos de Controle")

    st.info(
        "Objetivo: calcular rapidamente quantas camadas mais baixas do Windgram "
        "devem ser consideradas para um ponto escolhido no terreno, com base na "
        "distância desse ponto até o alvo."
    )

    # -----------------------------
    # 1. Dados automáticos
    # -----------------------------

    st.markdown("### 1. Dados automáticos")

    D_total_nm = float(st.session_state.get("distancia_velame_aberto_nm", 0.0))
    D_total_km = nm_para_km(D_total_nm)

    if "resumo_velame" in st.session_state and st.session_state.resumo_velame:
        camadas_totais = int(st.session_state.resumo_velame["qtd"])
    else:
        camadas_totais = 0

    c_auto1, c_auto2 = st.columns(2)

    with c_auto1:
        st.metric(
            "D total",
            f"{D_total_km:.3f} km",
            help="Distância total calculada na aba de Distância para Velame Aberto."
        )
        st.caption(f"Equivalente a {D_total_nm:.3f} NM")

    with c_auto2:
        st.metric(
            "Camadas totais",
            f"{camadas_totais}",
            help="Número de camadas consideradas a partir do Windgram processado."
        )

    if D_total_nm <= 0:
        st.warning(
            "Ainda não existe D calculado. Vá na aba 'Cálculo da Distância para Velame Aberto' "
            "e calcule a distância primeiro."
        )

    if camadas_totais <= 0:
        st.warning(
            "Ainda não existem camadas vindas do Windgram. Vá na aba 'Planejamento / Windgram', "
            "cole o Windgram e processe os dados."
        )

    st.divider()

    # -----------------------------
    # 2. Entrada do usuário
    # -----------------------------

    st.markdown("### 2. Distância do ponto de controle")

    distancia_pc_km = st.number_input(
        "Distância do ponto de controle até o alvo (km)",
        value=0.0,
        step=0.1,
        min_value=0.0,
        help="Informe a distância em linha reta do ponto escolhido no terreno até o alvo."
    )

    distancia_pc_nm = km_para_nm(distancia_pc_km)

    st.caption(f"Conversão automática: {distancia_pc_km:.3f} km = {distancia_pc_nm:.3f} NM")
    FS_kft = float(st.session_state.get("fs_velame_aberto_kft", 2.0))
    CteHz = float(st.session_state.get("ctehz_velame_aberto_kt", 23.3))
    K = float(st.session_state.get("k_velame_aberto_kft_h", 40.6))

    st.divider()

    # -----------------------------
    # 3. Resultado direto
    # -----------------------------

    # -----------------------------
    # 3. Resultado direto
    # -----------------------------

    if st.button("Calcular ponto de controle", type="primary"):
        if D_total_nm <= 0:
            st.error("Calcule primeiro o D na aba de Distância para Velame Aberto.")
        elif camadas_totais <= 0:
            st.error("Processe primeiro o Windgram na aba Planejamento / Windgram.")
        elif distancia_pc_km <= 0:
            st.error("Informe a distância do ponto de controle ao alvo.")
        elif distancia_pc_nm > D_total_nm:
            st.error("A distância do ponto não pode ser maior que o D total.")
        elif "df_windgram" not in st.session_state:
            st.error("Tabela Windgram não encontrada. Processe o Windgram novamente.")
        else:
            camadas_exatas = (distancia_pc_nm * camadas_totais) / D_total_nm
            camadas_consideradas = round(camadas_exatas)
            
            if camadas_consideradas < 1:
                camadas_consideradas = 1
            if camadas_consideradas > camadas_totais:
                camadas_consideradas = camadas_totais

            resumo_pc, df_camadas_pc = calcular_media_camadas_mais_baixas(
                st.session_state.df_windgram,
                camadas_consideradas
            )

            # Salva o cálculo na memória para não sumir ao clicar em outros botões
            st.session_state.calculo_pc_atual = {
                "camadas_consideradas": camadas_consideradas,
                "resumo_pc": resumo_pc,
                "df_camadas_pc": df_camadas_pc,
                "distancia_pc_nm": distancia_pc_nm,
                "distancia_pc_km": distancia_pc_km
            }

    # ==========================================
    # EXIBIÇÃO E ADIÇÃO NO BANCO DE DADOS
    # ==========================================
    
    # Só exibe os dados se o cálculo já foi feito e está na memória
    if "calculo_pc_atual" in st.session_state:
        dados_pc = st.session_state.calculo_pc_atual
        resumo_pc = dados_pc["resumo_pc"]
        
        r1, r2 = st.columns(2)
        with r1:
            st.metric("Camadas consideradas", f"{dados_pc['camadas_consideradas']}")
        with r2:
            st.metric("Média dos ventos nessas camadas", f"{resumo_pc['vento_medio']:.1f} kt")

        st.divider()

        st.markdown("### Dados para consulta")
        d1, d2, d3 = st.columns(3)
        with d1:
            st.metric("D usado", f"{dados_pc['distancia_pc_nm']:.3f} NM")
            st.caption(f"{dados_pc['distancia_pc_km']:.3f} km")
        with d2:
            st.metric("V usado", f"{resumo_pc['vento_medio']:.1f} kt")
        with d3:
            st.metric("Direção ponderada", f"{resumo_pc['direcao_ponderada']:.0f}°")

        d4, d5, d6 = st.columns(3)
        with d4:
            st.metric("FS", f"{FS_kft:.1f} kft")
        with d5:
            st.metric("CteHz", f"{CteHz:.1f} kt")
        with d6:
            st.metric("K", f"{K:.1f} kft/h")

        st.divider()
        
        st.markdown("### Cálculo da Altura P Ctle")
        v_usado = resumo_pc['vento_medio']
        vel_efetiva = CteHz + v_usado

        if vel_efetiva > 0:
            # O cálculo original te dá o valor em kft
            w_calculado_kft = ((dados_pc['distancia_pc_nm'] * K) / vel_efetiva) + FS_kft
            
            # Multiplicamos por 1000 para converter para ft
            w_calculado_ft = w_calculado_kft * 1000
            
            # Formatação para o padrão brasileiro (trocando vírgula por ponto e vice-versa)
            w_ft_formatado = f"{w_calculado_ft:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            
            st.metric(
                label="Altura P Ctle calculada",
                value=f"{w_ft_formatado} ft",
                help="Calculado a partir do D usado e média de ventos das camadas inferiores."
            )
            
            observacao_pc = st.text_input(
                "Observação / Referência Visual",
                placeholder="Ex: Curva de rio, estrada de terra...",
                key="observacao_manual_pc"
            )
            
            # Botão para salvar no banco
            if st.button("➕ Adicionar P Ctle", type="secondary"):
                novo_ponto = {
                    "ID": len(st.session_state.pontos_controle) + 1,
                    "Dist (km)": round(dados_pc['distancia_pc_km'], 3),
                    "Alt P Ctle (ft)": round(w_calculado_ft, 2), # <--- Agora salva em ft na tabela!
                    "Vento Médio (kt)": round(v_usado, 1),
                    "Camadas": dados_pc['camadas_consideradas'],
                    "Observação": observacao_pc
                }
                st.session_state.pontos_controle.append(novo_ponto)
                st.success(f"Ponto {novo_ponto['ID']} salvo com sucesso!")
                st.rerun() # Atualiza a tela instantaneamente para mostrar na tabela
                
        else:
            st.error("A velocidade efetiva (CteHz + V usado) é menor ou igual a zero.")
    # -----------------------------
    # TABELA DE REGISTROS (BANCO)
    # -----------------------------
    st.divider()
    st.markdown("### 🗂️ Pontos de Controle Registrados")

    if not st.session_state.pontos_controle:
        st.info("Nenhum ponto de controle adicionado ainda para este salto.")
    else:
        # Transforma a lista em DataFrame para ficar bonito no app
        df_banco = pd.DataFrame(st.session_state.pontos_controle)
        st.dataframe(df_banco, use_container_width=True, hide_index=True)
        
        # Opção para excluir dados se errar
        if st.button("🗑️ Limpar Todos os Pontos"):
            st.session_state.pontos_controle = []
            st.rerun()
    # -----------------------------
    # EXPORTAÇÃO KMZ (GERAL / VELAME ABERTO)
    # -----------------------------
    st.divider()
    st.markdown("### 🌍 Exportar Planejamento (Google Earth)")

    if st.session_state.get("df_windgram") is not None and st.session_state.get("distancia_velame_aberto_nm", 0) > 0:
        
        tipo_lancamento = st.radio(
            "Selecione o Tipo de Lançamento (Para Navegação):",
            ["Lançamento de Nariz", "Lançamento de Cauda"],
            horizontal=True,
            key="tipo_lanc_geral"
        )
        
# Parâmetros de Lançamento integrados (Aeronave, Inércia e Dispersão)
        st.markdown("#### Parâmetros da Aeronave e Dispersão")
        c_anv1, c_anv2 = st.columns(2)
        with c_anv1:
            aeronave_kmz = st.selectbox(
                "Aeronave",
                ["C-105 Amazonas / KC-390", "C-95 Bandeirante", "C-98 Caravan", "Outra (Manual)"],
                key="sel_aero_kmz"
            )
        with c_anv2:
            if aeronave_kmz == "C-105 Amazonas / KC-390":
                velocidade_anv_kmz = 70.0
                offset_inercia_kmz = 300.0
                st.number_input("Velocidade (m/s)", value=70.0, disabled=True, key="num_aero_kmz1")
            elif aeronave_kmz == "C-95 Bandeirante":
                velocidade_anv_kmz = 60.0
                offset_inercia_kmz = 150.0
                st.number_input("Velocidade (m/s)", value=60.0, disabled=True, key="num_aero_kmz2")
            elif aeronave_kmz == "C-98 Caravan":
                velocidade_anv_kmz = 45.0
                offset_inercia_kmz = 150.0
                st.number_input("Velocidade (m/s)", value=45.0, disabled=True, key="num_aero_kmz3")
            else:
                velocidade_anv_kmz = st.number_input("Velocidade (m/s)", value=70.0, step=1.0, key="num_aero_kmz4")
                offset_inercia_kmz = 300.0

        c_disp1, c_disp2 = st.columns(2)
        with c_disp1:
            num_blocos_kmz = st.number_input("Número de Blocos", min_value=1, value=1, step=1, key="nb_kmz_geral")
        with c_disp2:
            int_blocos_kmz = st.number_input("Intervalo (s)", min_value=0.0, value=1.0, step=0.5, key="ib_kmz_geral")

        dispersao_kmz_m = num_blocos_kmz * int_blocos_kmz * velocidade_anv_kmz
        deslocamento_total_m = offset_inercia_kmz + dispersao_kmz_m

        if tipo_lancamento == "Lançamento de Nariz":
            st.info(f"📍 O PS será recuado **{deslocamento_total_m:.0f} metros** a partir do Ponto D (Inércia: {offset_inercia_kmz}m + Dispersão: {dispersao_kmz_m:.0f}m).")
        elif tipo_lancamento == "Lançamento de Cauda":
            st.info("📍 O PS será **cravado no limite da Distância D**. (Arrasto e dispersão são desconsiderados no salto de cauda para Velame Aberto).")
        if st.button("Gerar Arquivo KMZ", type="primary", key="btn_kmz_geral"):
            registrar_log_missao("Gerou KMZ Velame Aberto") # <--- RASTREADOR
            import math
            
                        # Dados base
            lat_alvo = st.session_state.lat
            lon_alvo = st.session_state.lon
            dist_total_km = nm_para_km(st.session_state.distancia_velame_aberto_nm)

            # Direção verdadeira/geográfica para Google Earth / KMZ
            # Não aplica declinação magnética.
            azimute_vento = float(
            st.session_state.get(
        "direcao_vento_verdadeira_kmz",
        st.session_state.resumo_velame.get(
            "direcao_ponderada",
            st.session_state.resumo_velame.get("direcao_ponderada", 0.0)
        )
    )
)

            # Ponto final da reta vermelha (PS)
            lat_fim, lon_fim = calcular_coordenada_destino(
                lat_alvo, lon_alvo, dist_total_km, azimute_vento
            )

            # Detecta Dog Leg
            dog_leg = detectar_dog_leg(st.session_state.get("df_windgram"), limite_graus=90)

            dog_leg_geo = None
            if dog_leg:
                dog_leg_geo = resolver_geometria_dog_leg(
                    lat_alvo=lat_alvo,
                    lon_alvo=lon_alvo,
                    dist_total_km=dist_total_km,
                    azimute_vermelho=azimute_vento,
                    azimute_inferior=dog_leg["dir_inferior"],
                    azimute_superior=dog_leg["dir_superior"],
                    qtd_inferior=dog_leg["qtd_inferior"],
                    qtd_superior=dog_leg["qtd_superior"],
                )
           # Ponto final da reta vermelha (Distância D original)
            lat_d, lon_d = calcular_coordenada_destino(lat_alvo, lon_alvo, dist_total_km, azimute_vento)
            
            # =====================================================
            # CÁLCULO DA DIREÇÃO DE POUSO (2 CAMADAS MAIS BAIXAS)
            # =====================================================
            linha_pouso_kml = ""
            df_velame = st.session_state.get("df_windgram")
            if df_velame is not None and not df_velame.empty:
                base_velame = df_velame[df_velame["Fase"] == "Velame aberto"].copy()
                if not base_velame.empty:
                    # Ordena do mais baixo para o mais alto e pega os 2 primeiros
                    df_baixas = base_velame.sort_values("Altitude NOAA ft", ascending=True).head(2)
                    
                    dir_b = df_baixas["Direção °"].astype(float).tolist()
                    vel_b = df_baixas["Velocidade kt"].astype(float).tolist()
                    
                    if len(dir_b) > 0:
                        # Média Ponderada da Direção e Média Aritmética da Velocidade
                        dir_pouso = media_circular_ponderada(dir_b, vel_b)
                        if dir_pouso is None:
                            dir_pouso = sum(dir_b) / len(dir_b) # Fallback de segurança
                        
                        vel_pouso = sum(vel_b) / len(vel_b)
                        
                        # Calcula a ponta da linha a 0.5 km (500 metros) do alvo
                        lat_pouso, lon_pouso = calcular_coordenada_destino(lat_alvo, lon_alvo, 0.5, dir_pouso)
                        
                        linha_pouso_kml = f"""
    <Placemark>
      <name>Eixo de Pouso (500m)</name>
      <styleUrl>#linhaVerde</styleUrl>
      <LineString>
        <extrude>1</extrude>
        <tessellate>1</tessellate>
        <coordinates>
          {lon_alvo},{lat_alvo},0
          {lon_pouso},{lat_pouso},0
        </coordinates>
      </LineString>
    </Placemark>
    <Placemark>
      <name>Provável direção de pouso - {vel_pouso:.1f} kt</name>
      <description>Direção ponderada: {dir_pouso:.0f}° | Média calculada das {len(dir_b)} camada(s) mais baixa(s) do Windgram</description>
      <styleUrl>#iconePonto</styleUrl>
      <Point>
        <coordinates>{lon_pouso},{lat_pouso},0</coordinates>
      </Point>
    </Placemark>
"""

            # Foi adicionada a "linhaVerde" na tabela de estilos
            estilos_kml = """
            <Style id="linhaVermelha"><LineStyle><color>ff0000ff</color><width>4</width></LineStyle></Style>
            <Style id="linhaAmarela"><LineStyle><color>ff00ffff</color><width>4</width></LineStyle></Style>
            <Style id="linhaVerde"><LineStyle><color>ff00ff00</color><width>4</width></LineStyle></Style>
            <Style id="iconeAlvo"><IconStyle><Icon><href>http://maps.google.com/mapfiles/kml/paddle/red-stars.png</href></Icon></IconStyle></Style>
            <Style id="iconePonto"><IconStyle><Icon><href>http://maps.google.com/mapfiles/kml/paddle/ylw-blank.png</href></Icon></IconStyle></Style>
            """
            
            kml_str = f"""<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <name>Planejamento SLOP</name>
    {estilos_kml}

    <Placemark>
      <name>Alvo</name>
      <styleUrl>#iconeAlvo</styleUrl>
      <Point>
        <coordinates>{lon_alvo},{lat_alvo},0</coordinates>
      </Point>
    </Placemark>
    
    {linha_pouso_kml}

    <Placemark>
      <name>Distância D ({(dist_total_km * 1000):.0f} m)</name>
      <styleUrl>#linhaVermelha</styleUrl>
      <LineString>
        <extrude>1</extrude>
        <tessellate>1</tessellate>
        <coordinates>
          {lon_alvo},{lat_alvo},0
          {lon_d},{lat_d},0
        </coordinates>
      </LineString>
    </Placemark>
"""
            
            if tipo_lancamento == "Lançamento de Nariz":
                # FÍSICA CORRIGIDA: O avião voa do Alvo para o D. 
                # Lança ANTES de chegar no D (voltamos pelo contra_azimute para ficar mais perto do alvo)
                lat_ps_final, lon_ps_final = calcular_coordenada_destino(
                    lat_d, lon_d, deslocamento_total_m / 1000.0, contra_azimute(azimute_vento)
                )
                st.session_state.ps_origem = "PONTO DE SAÍDA - Nariz (Compensado)"
                
                kml_str += f"""
    <Placemark>
      <name>Inércia + Dispersão Nariz ({deslocamento_total_m:.0f} m)</name>
      <styleUrl>#linhaAzul</styleUrl>
      <LineString>
        <extrude>1</extrude>
        <tessellate>1</tessellate>
        <coordinates>
          {lon_d},{lat_d},0
          {lon_ps_final},{lat_ps_final},0
        </coordinates>
      </LineString>
    </Placemark>
"""

            elif tipo_lancamento == "Lançamento de Cauda":
                # O avião vem de trás, passa pelo D e continua voando. 
                # Dispersão e Inércia irrelevantes no Velame Aberto: PS fica cravado no limite D.
                lat_ps_final = lat_d
                lon_ps_final = lon_d
                st.session_state.ps_origem = "PONTO DE SAÍDA - Cauda (Cravado no D)"
                
                if dog_leg:
                    st.info(f"Dog Leg detectado. Exportando KMZ com a quebra.")

            st.session_state.ps_lat = lat_ps_final
            st.session_state.ps_lon = lon_ps_final

            # === MÁGICA: CÁLCULO E TRAÇADO DO 1' FORA ===
            dist_1_fora_km = (velocidade_anv_kmz * 60) / 1000.0
            if tipo_lancamento == "Lançamento de Nariz":
                az_1_fora = contra_azimute(azimute_vento) # Correção: volta a favor do vento
            else:
                az_1_fora = azimute_vento # Correção: volta contra o vento                
            lat_1_fora, lon_1_fora = calcular_coordenada_destino(lat_ps_final, lon_ps_final, dist_1_fora_km, az_1_fora)

            kml_str += f"""
    <Placemark>
      <name>PONTO DE SAÍDA - PS</name>
      <description>Início da Navegação</description>
      <styleUrl>#iconePonto</styleUrl>
      <Point>
        <coordinates>{lon_ps_final},{lat_ps_final},0</coordinates>
      </Point>
    </Placemark>
    
    <Placemark>
      <name>Trajeto 1' FORA</name>
      <styleUrl>#linhaAmarela</styleUrl>
      <LineString>
        <extrude>1</extrude>
        <tessellate>1</tessellate>
        <coordinates>
          {lon_ps_final},{lat_ps_final},0
          {lon_1_fora},{lat_1_fora},0
        </coordinates>
      </LineString>
    </Placemark>

    <Placemark>
      <name>1' FORA</name>
      <description>Aeronave a 1 minuto do Ponto de Saída</description>
      <styleUrl>#iconePonto</styleUrl>
      <Point>
        <coordinates>{lon_1_fora},{lat_1_fora},0</coordinates>
      </Point>
    </Placemark>
"""
                        # Se houver Dog Leg, desenha as 2 retas azuis
            if dog_leg and dog_leg_geo:
                lat_quebra = dog_leg_geo["lat_quebra"]
                lon_quebra = dog_leg_geo["lon_quebra"]

                dist_inf = dog_leg_geo["dist_inferior_km"]
                dist_sup = dog_leg_geo["dist_superior_km"]

                dir_inf = dog_leg["dir_inferior"]
                dir_sup = dog_leg["dir_superior"]

                kml_str += f"""
    <Placemark>
      <name>Dog Leg - Trecho Inferior ({dist_inf:.2f} km)</name>
      <description>Direção média do bloco inferior: {dir_inf:.0f}°</description>
      <styleUrl>#linhaDogLeg</styleUrl>
      <LineString>
        <extrude>1</extrude>
        <tessellate>1</tessellate>
        <coordinates>
          {lon_alvo},{lat_alvo},0
          {lon_quebra},{lat_quebra},0
        </coordinates>
      </LineString>
    </Placemark>

    <Placemark>
      <name>Dog Leg - Trecho Superior ({dist_sup:.2f} km)</name>
      <description>Direção média do bloco superior: {dir_sup:.0f}°</description>
      <styleUrl>#linhaDogLeg</styleUrl>
      <LineString>
        <extrude>1</extrude>
        <tessellate>1</tessellate>
        <coordinates>
          {lon_quebra},{lat_quebra},0
          {lon_fim},{lat_fim},0
        </coordinates>
      </LineString>
    </Placemark>

    <Placemark>
      <name>Ponto de Quebra do Dog Leg</name>
      <description>Transição entre os blocos de vento</description>
      <styleUrl>#iconePonto</styleUrl>
      <Point>
        <coordinates>{lon_quebra},{lat_quebra},0</coordinates>
      </Point>
    </Placemark>
"""
            # Marcadores dos Pontos de Controle
            for p in st.session_state.pontos_controle:
                lat_pc, lon_pc = calcular_coordenada_destino(lat_alvo, lon_alvo, p["Dist (km)"], azimute_vento)
                
                altura_kft = p['Alt P Ctle (ft)'] / 1000
                altura_formatada = f"{altura_kft:.1f}".replace(".", ",")
                nome_pc = f"P Ct {p['ID']} - {altura_formatada}"
                
                obs = p.get('Observação', '')
                
                kml_str += f"""
    <Placemark>
      <name>{nome_pc}</name>
      <description>{obs}</description>
      <styleUrl>#iconePonto</styleUrl>
      <Point>
        <coordinates>{lon_pc},{lat_pc},0</coordinates>
      </Point>
    </Placemark>
"""
            
            kml_str += """  </Document>\n</kml>"""
            
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                zip_file.writestr("planejamento.kml", kml_str.encode("utf-8"))
            
            kmz_bytes = zip_buffer.getvalue()
            
            st.success("KMZ gerado com sucesso!")
            st.download_button(
                label="🗺️ Baixar KMZ para Google Earth",
                data=kmz_bytes,
                file_name="Planejamento_SLOP.kmz",
                mime="application/vnd.google-earth.kmz"
            )
    else:
        st.info("Realize o cálculo da Distância Total e do Windgram nas abas anteriores para liberar a exportação.")
        # =====================================================
# ABA SALTO SOBRE O ALVO (DKVA)
# =====================================================

with aba_dkva:
    st.subheader("🎯 Planejamento para Saltos sobre o Alvo (DKVA)")
    st.info("Objetivo: Calcular a distância de lançamento (D) em metros usando a constante K=25.")

    # Verifica se já temos o vento calculado na primeira aba
    if "resumo_velame" not in st.session_state or not st.session_state.resumo_velame:
        st.warning("⚠️ Volte na aba 'Planejamento / Windgram' e processe o Windgram para obter o Vento Médio e a Direção.")
    else:
        v_usado = st.session_state.resumo_velame["vento_medio"]
        direcao_vento = float(
    st.session_state.get(
        "direcao_vento_verdadeira_kmz",
        st.session_state.resumo_velame.get(
            "direcao_ponderada",
            st.session_state.resumo_velame.get("direcao_ponderada", 0.0)
        )
    )
)
        K_alvo = 25

        st.markdown("### 1. Parâmetros")
        
        c_alvo1, c_alvo2, c_alvo3 = st.columns(3)
        
        with c_alvo1:
            st.metric("K (Constante)", f"{K_alvo}")
            
        with c_alvo2:
            st.metric("V (Vento Médio)", f"{v_usado:.1f} kt")
            
        with c_alvo3:
            # Puxa a altura da aba de planejamento e converte para kft automaticamente
            altura_comandamento_kft_dkva = float(st.session_state.get("altura_comandamento_ft", 1200.0)) / 1000.0
            
            A_kft = st.number_input(
                "A (Altura de abertura em kft)",
                value=altura_comandamento_kft_dkva,
                step=0.1,
                help="Ex: 1.2 significa 1.200 pés. (Sincronizado automaticamente com a Aba Planejamento)",
                key=f"dkva_a_auto_{int(st.session_state.get('altura_comandamento_ft', 1200.0))}"
            )
            # Salva na memória original para garantir que o Relatório TXT continue funcionando perfeitamente
            st.session_state.dkva_a_kft = A_kft

        st.divider()
        st.markdown("### 2. Resultado")

        # Cálculo da fórmula D = K * V * A (Resultado em metros)
        d_metros = K_alvo * v_usado * A_kft
        d_km = d_metros / 1000

        r1, r2, r3 = st.columns(3)
        with r1:
            st.metric("Distância (D)", f"{d_metros:.0f} metros")
        with r2:
            st.metric("Direção do Vento", f"{direcao_vento:.0f}°")
        with r3:
            st.metric("Eixo de Navegação (Nariz)", f"{contra_azimute(direcao_vento):.0f}°")

# =====================================================
        # EXPORTAÇÃO KMZ (DKVA TÁTICO)
        # =====================================================
        st.divider()
        st.markdown("### 🌍 Exportar Planejamento DKVA (Google Earth)")

        tipo_lanc_dkva = st.radio(
            "Selecione o Tipo de Lançamento:",
            ["Lançamento de Nariz", "Lançamento de Cauda", "Lançamento Boca do Cone"],
            horizontal=True,
            key="tipo_lanc_dkva_radio" # <--- ADICIONADO ISSO AQUI
        )

        offset_base = 0.0
        dispersao_m = 0.0

        if tipo_lanc_dkva in ["Lançamento de Cauda", "Lançamento Boca do Cone"]:
            st.markdown(f"#### Parâmetros para {tipo_lanc_dkva}")
            c_dk1, c_dk2, c_dk3 = st.columns(3)
            with c_dk1:
                num_blocos = st.number_input("Número de Blocos", min_value=1, value=1, step=1, key="nb_dkva")
            with c_dk2:
                int_blocos = st.number_input("Intervalo (s)", min_value=0.0, value=1.0, step=0.5, key="ib_dkva")
            with c_dk3:
                # Dicionário com [Velocidade (m/s), Offset Base (m)]
                dic_aero_dkva = {
                    "C 105 Amazonas": [70, 300],
                    "KC 390": [70, 300],
                    "Beech 99": [60, 150],
                    "Caravan": [40, 150],
                    "Gran Caravan": [40, 150],
                    "Helicóptero": [40, 150]
                }
                aero_esc = st.selectbox("Aeronave", list(dic_aero_dkva.keys()), key="ae_dkva")
                vel_aeronave = dic_aero_dkva[aero_esc][0]
                offset_base = dic_aero_dkva[aero_esc][1]

            dispersao_m = num_blocos * int_blocos * vel_aeronave
            if tipo_lanc_dkva == "Lançamento de Cauda":
                st.info("📍 O PS será **cravado no limite da Distância D**. (Arrasto e dispersão são desconsiderados no salto de cauda).")
            elif tipo_lanc_dkva == "Lançamento Boca do Cone":
                st.info(f"📏 A reta perpendicular de dispersão terá **{dispersao_m:.0f} metros** de comprimento total.")

        if st.button("🗺️ Gerar Arquivo KMZ do DKVA", type="primary"):
            registrar_log_missao("Gerou KMZ Salto DKVA") # <--- RASTREADOR
            import math

            lat_alvo = st.session_state.lat
            lon_alvo = st.session_state.lon
            lat_pl, lon_pl = calcular_coordenada_destino(lat_alvo, lon_alvo, d_km, direcao_vento)
            st.session_state.ps_lat = lat_pl
            st.session_state.ps_lon = lon_pl
            st.session_state.ps_origem = "PS DKVA"
# =====================================================
            # CÁLCULO DE ALTITUDE E DAA (QFE) DO PS
            # =====================================================
            st.markdown("#### 📍 Dados Atmosféricos do Ponto de Saída (PS)")
            
            with st.spinner("A consultar a base de dados meteorológica (QFE e Altitude)..."):
                alt_ps_ft = 0.0
                qfe_hpa = 0.0
                
                try:
                    # 1. Busca a Elevação Exata pura (em metros) e converte para pés
                    url_elev = f"https://api.open-meteo.com/v1/elevation?latitude={lat_pl}&longitude={lon_pl}"
                    resp_elev = requests.get(url_elev, timeout=5).json()
                    
                    if "elevation" in resp_elev and len(resp_elev["elevation"]) > 0:
                        alt_ps_m = float(resp_elev["elevation"][0])
                        alt_ps_ft = alt_ps_m * 3.28084
                    
                    # 2. Busca a Pressão Superficial (QFE) exata do Ponto de Saída
                    url_clima = f"https://api.open-meteo.com/v1/forecast?latitude={lat_pl}&longitude={lon_pl}&current=surface_pressure"
                    resp_clima = requests.get(url_clima, timeout=5).json()
                    
                    if "current" in resp_clima and "surface_pressure" in resp_clima["current"]:
                        qfe_hpa = float(resp_clima["current"]["surface_pressure"])
                        
                except Exception:
                    pass
            
            # --- MÁGICA DA INTEGRAÇÃO COM O FOLDER DO PILOTO ---
            # Salva os dados recém-calculados do DKVA nas mesmas variáveis globais que o Folder lê!
            if alt_ps_ft > 0:
                st.session_state.altitude_consulta_ft = alt_ps_ft
            if qfe_hpa > 0:
                st.session_state.qfe_consulta_hpa = qfe_hpa
            # ---------------------------------------------------

            c_ps1, c_ps2 = st.columns(2)
            
            with c_ps1:
                # Se falhar, deixa a caixa vazia para poder digitar manualmente
                val_alt = f"{alt_ps_ft:.0f} ft" if alt_ps_ft > 0 else ""
                st.text_input("Altitude do PS (MSL):", value=val_alt, key="ps_alt_input")
                
            with c_ps2:
                # Campo textual escrito DAA exibindo a resposta do QFE
                val_qfe = f"{qfe_hpa:.1f} hPa" if qfe_hpa > 0 else ""
                st.text_input("DAA :", value=val_qfe, key="ps_daa_input")
                
            st.divider()
            # ---> METROS CALCULADOS AQUI E DENTRO DO BOTÃO <---
            d_metros = d_km * 1000.0

            # Funções matemáticas blindadas para traçar os pontos no mapa
            def calc_rumo(lat1, lon1, lat2, lon2):
                l1, ln1, l2, ln2 = map(math.radians, [lat1, lon1, lat2, lon2])
                d_lon = ln2 - ln1
                x = math.sin(d_lon) * math.cos(l2)
                y = math.cos(l1) * math.sin(l2) - (math.sin(l1) * math.cos(l2) * math.cos(d_lon))
                return (math.degrees(math.atan2(x, y)) + 360) % 360

            def mover_ponto(lat, lon, dist_km, azimute_graus):
                R = 6371.0
                l_rad, ln_rad, az_rad = map(math.radians, [lat, lon, azimute_graus])
                lat_n = math.asin(math.sin(l_rad) * math.cos(dist_km/R) + math.cos(l_rad) * math.sin(dist_km/R) * math.cos(az_rad))
                lon_n = ln_rad + math.atan2(math.sin(az_rad) * math.sin(dist_km/R) * math.cos(l_rad), math.cos(dist_km/R) - math.sin(l_rad) * math.sin(lat_n))
                return math.degrees(lat_n), math.degrees(lon_n)

            # O vento sopra do Alvo para o Ponto D
            az_vento = calc_rumo(lat_alvo, lon_alvo, lat_pl, lon_pl)
            dec_mag = st.session_state.get("declinacao", 0.0)
            
            # =====================================================
            # CÁLCULO DA DIREÇÃO DE POUSO (2 CAMADAS MAIS BAIXAS)
            # =====================================================
            linha_pouso_kml = ""
            df_velame = st.session_state.get("df_windgram")
            if df_velame is not None and not df_velame.empty:
                base_velame = df_velame[df_velame["Fase"] == "Velame aberto"].copy()
                if not base_velame.empty:
                    df_baixas = base_velame.sort_values("Altitude NOAA ft", ascending=True).head(2)
                    dir_b = df_baixas["Direção °"].astype(float).tolist()
                    vel_b = df_baixas["Velocidade kt"].astype(float).tolist()
                    
                    if len(dir_b) > 0:
                        dir_pouso = media_circular_ponderada(dir_b, vel_b)
                        if dir_pouso is None:
                            dir_pouso = sum(dir_b) / len(dir_b)
                        vel_pouso = sum(vel_b) / len(vel_b)
                        
                        lat_pouso, lon_pouso = mover_ponto(lat_alvo, lon_alvo, 0.5, dir_pouso)
                        
                        linha_pouso_kml = f"""
            <Placemark>
              <name>Eixo de Pouso (500m)</name>
              <styleUrl>#linhaVerde</styleUrl>
              <LineString>
                <extrude>1</extrude>
                <tessellate>1</tessellate>
                <coordinates>
                  {lon_alvo},{lat_alvo},0
                  {lon_pouso},{lat_pouso},0
                </coordinates>
              </LineString>
            </Placemark>
            <Placemark>
              <name>Provável direção de pouso - {vel_pouso:.1f} kt</name>
              <description>Direção ponderada: {dir_pouso:.0f}° | Média das {len(dir_b)} camadas mais baixas</description>
              <styleUrl>#iconePonto</styleUrl>
              <Point>
                <coordinates>{lon_pouso},{lat_pouso},0</coordinates>
              </Point>
            </Placemark>
            """

            # Estilos visuais do Google Earth
            estilos_kml = """
            <Style id="linhaVermelha"><LineStyle><color>ff0000ff</color><width>4</width></LineStyle></Style>
            <Style id="linhaAzul"><LineStyle><color>ffff0000</color><width>4</width></LineStyle></Style>
            <Style id="linhaAmarela"><LineStyle><color>ff00ffff</color><width>4</width></LineStyle></Style>
            <Style id="linhaVerde"><LineStyle><color>ff00ff00</color><width>4</width></LineStyle></Style>
            <Style id="iconeAlvo"><IconStyle><Icon><href>http://maps.google.com/mapfiles/kml/paddle/red-stars.png</href></Icon></IconStyle></Style>
            <Style id="iconePonto"><IconStyle><Icon><href>http://maps.google.com/mapfiles/kml/paddle/ylw-blank.png</href></Icon></IconStyle></Style>
            """
            
            # Elementos Fixos: Alvo, Reta D (Sempre Vermelha na extensão D) e Linha Verde de Pouso
            kml_elementos = f"""
            <Placemark>
              <name>Alvo</name>
              <styleUrl>#iconeAlvo</styleUrl>
              <Point><coordinates>{lon_alvo},{lat_alvo},0</coordinates></Point>
            </Placemark>
            
            {linha_pouso_kml}

            <Placemark>
              <name>Distância D (Limite: {d_metros:.0f} m)</name>
              <styleUrl>#linhaVermelha</styleUrl>
              <LineString><extrude>1</extrude><tessellate>1</tessellate>
                <coordinates>{lon_alvo},{lat_alvo},0 {lon_pl},{lat_pl},0</coordinates>
              </LineString>
            </Placemark>
            <Placemark>
              <name>Ponto Limite D</name>
              <styleUrl>#iconePonto</styleUrl>
              <Point><coordinates>{lon_pl},{lat_pl},0</coordinates></Point>
            </Placemark>
            """

            if tipo_lanc_dkva == "Lançamento de Nariz":
                # O arrasto de 300m vai do alvo para TRÁS (Contra azimute da direção de D)
                az_contra_d = (az_vento + 180) % 360
                lat_ps_nariz, lon_ps_nariz = mover_ponto(lat_alvo, lon_alvo, 300.0 / 1000.0, az_contra_d)
                
                st.session_state.ps_lat = lat_ps_nariz
                st.session_state.ps_lon = lon_ps_nariz
                st.session_state.ps_origem = "PS DKVA (Nariz - 300m Contra D)"

                kml_elementos += f"""
                <Placemark>
                  <name>Inércia Aeronave (300 m)</name>
                  <styleUrl>#linhaAzul</styleUrl>
                  <LineString><extrude>1</extrude><tessellate>1</tessellate>
                    <coordinates>{lon_alvo},{lat_alvo},0 {lon_ps_nariz},{lat_ps_nariz},0</coordinates>
                  </LineString>
                </Placemark>
                <Placemark>
                  <name>PS (Ponto de Saída)</name>
                  <description>Altura: {A_kft:.1f} kft | Vento: {v_usado:.1f} kt | Luz verde a 300m do Alvo (Contra Azimute).</description>
                  <styleUrl>#iconePonto</styleUrl>
                  <Point><coordinates>{lon_ps_nariz},{lat_ps_nariz},0</coordinates></Point>
                </Placemark>
                """

            elif tipo_lanc_dkva == "Lançamento de Cauda":
                # Dispersão e Inércia irrelevantes no Lançamento de Cauda: PS fica cravado no limite D.
                lat_ps_cauda = lat_pl
                lon_ps_cauda = lon_pl
                
                st.session_state.ps_lat = lat_ps_cauda
                st.session_state.ps_lon = lon_ps_cauda
                st.session_state.ps_origem = "PS DKVA (Cauda - Cravado no D)"
                
                kml_elementos += f"""
                <Placemark>
                  <name>PS Final (Cauda - Cravado no D)</name>
                  <description>Arrasto e dispersão desconsiderados.</description>
                  <styleUrl>#iconePonto</styleUrl>
                  <Point><coordinates>{lon_ps_cauda},{lat_ps_cauda},0</coordinates></Point>
                </Placemark>
                """

            elif tipo_lanc_dkva == "Lançamento Boca do Cone":
                # Traça 90º para a direita e esquerda em relação ao eixo do vento a partir da Distância D
                az_direita = (az_vento + 90) % 360
                az_esquerda = (az_vento - 90) % 360
                
                # Desconta a declinação magnética para exibir os rumos puros
                az_mag_dir = (az_direita - dec_mag) % 360
                az_mag_esq = (az_esquerda - dec_mag) % 360
                
                meia_disp_km = (dispersao_m / 2.0) / 1000.0
                
                # Desenha os dois pontos nas extremidades da reta perpendicular
                lat_dir, lon_dir = mover_ponto(lat_pl, lon_pl, meia_disp_km, az_direita)
                lat_esq, lon_esq = mover_ponto(lat_pl, lon_pl, meia_disp_km, az_esquerda)
                
                kml_elementos += f"""
                <Placemark>
                  <name>Dispersão Perpendicular (Boca do Cone - {dispersao_m:.0f}m)</name>
                  <styleUrl>#linhaAmarela</styleUrl>
                  <LineString><extrude>1</extrude><tessellate>1</tessellate>
                    <coordinates>{lon_esq},{lat_esq},0 {lon_dir},{lat_dir},0</coordinates>
                  </LineString>
                </Placemark>
                <Placemark>
                  <name>PS Direita (Rumo Mag: {az_mag_esq:.0f}°)</name>
                  <styleUrl>#iconePonto</styleUrl>
                  <Point><coordinates>{lon_dir},{lat_dir},0</coordinates></Point>
                </Placemark>
                <Placemark>
                  <name>PS Esquerda (Rumo Mag: {az_mag_dir:.0f}°)</name>
                  <styleUrl>#iconePonto</styleUrl>
                  <Point><coordinates>{lon_esq},{lat_esq},0</coordinates></Point>
                </Placemark>
                """

            # Monta e comprime o KML
            kml_str = f'<?xml version="1.0" encoding="UTF-8"?>\n<kml xmlns="http://www.opengis.net/kml/2.2">\n  <Document>\n    <name>Planejamento DKVA Tático</name>\n    {estilos_kml}\n    {kml_elementos}\n  </Document>\n</kml>'
            zip_buffer_dkva = io.BytesIO()
            with zipfile.ZipFile(zip_buffer_dkva, "w", zipfile.ZIP_DEFLATED) as zip_file:
                zip_file.writestr("planejamento_dkva_tatico.kml", kml_str.encode("utf-8"))
                
            kmz_bytes = zip_buffer_dkva.getvalue()
            
            st.download_button(
                label="🗺️ Baixar KMZ Tático (DKVA)",
                data=kmz_bytes,
                file_name="DKVA_Tatico.kmz",
                mime="application/vnd.google-earth.kmz"
            )

    # =====================================================
    # CÁLCULO DE AJUSTE ALTIMÉTRICO (DKVA)
    # =====================================================
    st.divider()
    st.markdown("### 🛫 Aeródromo de Partida e Ajuste Altimétrico")

    # --- AERÓDROMOS PRÉ-CADASTRADOS ---
    aeros_cadastrados_dkva = {
        "Personalizado / Outro": None,
        "Aeródromo de Goiânia": {"lat": -16.630929, "lon": -49.217527},
        "Skydive Cerrado": {"lat": -16.362042, "lon": -48.928371}
    }
    
    aero_selecionado_dkva = st.selectbox(
        "Aeródromos Conhecidos", 
        list(aeros_cadastrados_dkva.keys()),
        help="Selecione um aeródromo salvo para preencher as coordenadas e o mapa automaticamente.",
        key="sel_aero_conhecido_dkva"
    )
    
    if aero_selecionado_dkva != "Personalizado / Outro" and st.session_state.get("ultimo_aero_selecionado_dkva") != aero_selecionado_dkva:
        st.session_state.lat_aerodromo_partida = aeros_cadastrados_dkva[aero_selecionado_dkva]["lat"]
        st.session_state.lon_aerodromo_partida = aeros_cadastrados_dkva[aero_selecionado_dkva]["lon"]
        st.session_state.mapa_aerodromo_rev += 1
        st.session_state.ultimo_aero_selecionado_dkva = aero_selecionado_dkva
        st.rerun()
    elif aero_selecionado_dkva == "Personalizado / Outro":
        st.session_state.ultimo_aero_selecionado_dkva = "Personalizado / Outro"
    # ----------------------------------------

    with st.form("form_busca_aerodromo_dkva"):
        busca_aero_dkva = st.text_input(
            "Buscar aeródromo/localidade de partida",
            placeholder="Ex: Campo Grande MS, Anápolis, Goiânia..."
        )
        btn_busca_dkva = st.form_submit_button("🔎 Buscar local do aeródromo")

    if btn_busca_dkva and busca_aero_dkva.strip():
        res_aero = buscar_localidade(busca_aero_dkva)
        if res_aero:
            st.session_state.lat_aerodromo_partida = float(res_aero["lat"])
            st.session_state.lon_aerodromo_partida = float(res_aero["lon"])
            st.session_state.mapa_aerodromo_rev += 1
            st.success(f"{res_aero['nome']} | Fonte: {res_aero.get('fonte', 'Busca online')}")
            st.rerun()
        else:
            st.error("Aeródromo/localidade não encontrado.")

    c_aero_lat_dkva, c_aero_lon_dkva = st.columns(2)
    with c_aero_lat_dkva:
        lat_aero_dkva = st.number_input(
            "Latitude do aeródromo de partida",
            value=float(st.session_state.lat_aerodromo_partida),
            step=0.0001,
            format="%.6f",
            key=f"lat_aero_dkva_input_{st.session_state.mapa_aerodromo_rev}"
        )
    with c_aero_lon_dkva:
        lon_aero_dkva = st.number_input(
            "Longitude do aeródromo de partida",
            value=float(st.session_state.lon_aerodromo_partida),
            step=0.0001,
            format="%.6f",
            key=f"lon_aero_dkva_input_{st.session_state.mapa_aerodromo_rev}"
        )

    st.session_state.lat_aerodromo_partida = lat_aero_dkva
    st.session_state.lon_aerodromo_partida = lon_aero_dkva

    # --- MAPA INTERATIVO DO AERÓDROMO NO DKVA ---
    st.markdown("###### Selecionar aeródromo no mapa")

    mapa_aero_dkva = criar_mapa_base(
        [
            st.session_state.lat_aerodromo_partida,
            st.session_state.lon_aerodromo_partida
        ],
        zoom=12
    )

    folium.Marker(
        location=[
            st.session_state.lat_aerodromo_partida,
            st.session_state.lon_aerodromo_partida
        ],
        popup="Aeródromo de partida",
        tooltip="Aeródromo de partida",
        icon=folium.Icon(color="blue", icon="flag"),
    ).add_to(mapa_aero_dkva)

    resultado_mapa_dkva = st_folium(
        mapa_aero_dkva,
        width=None,
        height=360,
        key=f"mapa_aero_dkva_widget_{st.session_state.mapa_aerodromo_rev}",
        returned_objects=["last_clicked"],
    )

    if resultado_mapa_dkva and resultado_mapa_dkva.get("last_clicked"):
        lat_click_dkva = float(resultado_mapa_dkva["last_clicked"]["lat"])
        lon_click_dkva = float(resultado_mapa_dkva["last_clicked"]["lng"])

        novo_clique_dkva = [round(lat_click_dkva, 7), round(lon_click_dkva, 7)]

        if st.session_state.get("ultimo_clique_aero_dkva") != novo_clique_dkva:
            st.session_state.ultimo_clique_aero_dkva = novo_clique_dkva
            st.session_state.lat_aerodromo_partida = lat_click_dkva
            st.session_state.lon_aerodromo_partida = lon_click_dkva
            st.session_state.mapa_aerodromo_rev += 1
            st.rerun()

    st.caption(
        f"Aeródromo selecionado: "
        f"{st.session_state.lat_aerodromo_partida:.6f}, "
        f"{st.session_state.lon_aerodromo_partida:.6f}"
    )

    if st.button("🧮 Calcular Diferença Altimétrica", key="btn_calc_alt_dkva"):
        alt_aero_ft, _ = consultar_terreno_e_pressao(lat_aero_dkva, lon_aero_dkva)
        alt_alvo_ft = float(st.session_state.get("altitude_ft", 0.0))
        
        st.session_state.altitude_aerodromo_partida_ft = alt_aero_ft
        if alt_aero_ft is not None:
            st.session_state.altimetro_aerodromo_alvo_ft = alt_aero_ft - alt_alvo_ft
        else:
            st.session_state.altimetro_aerodromo_alvo_ft = None

    c_res1, c_res2 = st.columns(2)
    with c_res1:
        alt_aero_val = st.session_state.get("altitude_aerodromo_partida_ft")
        if alt_aero_val is not None:
            st.metric("Altitude do aeródromo", f"{alt_aero_val:,.0f} ft".replace(",", "X").replace(".", ",").replace("X", "."))
        else:
            st.metric("Altitude do aeródromo", "—")
            
    with c_res2:
        alt_dif_val = st.session_state.get("altimetro_aerodromo_alvo_ft")
        if alt_dif_val is not None:
            st.metric("Diferença Altimétrica", f"{alt_dif_val:,.0f} ft".replace(",", "X").replace(".", ",").replace("X", "."))
            st.caption("Altitude do aeródromo de partida - altitude do alvo")
        else:
            st.metric("Diferença Altimétrica", "—")
            # =====================================================
# ABA FOLDER DO PILOTO
# =====================================================

with aba_folder:
    st.subheader("📄 Folder do Piloto")

    if not DOCX_OK:
        st.error(
            "A biblioteca python-docx não está instalada. "
            "Adicione 'python-docx' ao requirements.txt."
        )
    else:
        st.info(
            "Esta aba gera um arquivo Word com os dados já calculados no app."
        )

        st.markdown("### Parâmetros da Aeronave")
        c_anv_f1, c_anv_f2 = st.columns(2)
        with c_anv_f1:
            aeronave_folder = st.selectbox(
                "Aeronave",
                ["C-105 Amazonas / KC-390", "C-95 Bandeirante", "C-98 Caravan", "Outra (Manual)"],
                key="sel_aero_folder"
            )
        with c_anv_f2:
            if aeronave_folder == "C-105 Amazonas / KC-390":
                velocidade_anv = 70.0
                st.number_input("Velocidade (m/s)", value=70.0, disabled=True, key="num_aero_f1")
            elif aeronave_folder == "C-95 Bandeirante":
                velocidade_anv = 60.0
                st.number_input("Velocidade (m/s)", value=60.0, disabled=True, key="num_aero_f2")
            elif aeronave_folder == "C-98 Caravan":
                velocidade_anv = 45.0
                st.number_input("Velocidade (m/s)", value=45.0, disabled=True, key="num_aero_f3")
            else:
                velocidade_anv = st.number_input("Velocidade (m/s)", value=70.0, step=1.0, key="num_aero_f4")
        distancia_teorica_1_fora_m = velocidade_anv * 60

        st.caption(
            f"Distância teórica 1' fora: "
            f"{distancia_teorica_1_fora_m:,.0f} m".replace(",", ".")
        )
        tipo_lancamento_folder = st.selectbox(
            "Tipo de lançamento para o folder",
            ["Lançamento de Nariz", "Lançamento de Cauda"],
            key="folder_tipo_lancamento"
        )
        st.markdown("### Referências 4' fora / 1' fora / PS")

        # --- CÁLCULO DO 1' FORA AUTOMÁTICO COM BLINDAGEM DE SESSÃO ---
        lat_1_fora_calc, lon_1_fora_calc = None, None
        
        if "ps_lat" in st.session_state and "ps_lon" in st.session_state:
            az_vento_verdadeiro = float(st.session_state.get("direcao_vento_verdadeira_kmz", 0.0))
            
            # Avião voa do 1' fora para o PS. Retrocedemos do PS para achar o 1' fora.
            if tipo_lancamento_folder == "Lançamento de Nariz":
                az_para_1_fora = contra_azimute(az_vento_verdadeiro) # Correção aplicada
            else:
                az_para_1_fora = az_vento_verdadeiro # Correção aplicada
                
            lat_1_fora_calc, lon_1_fora_calc = calcular_coordenada_destino(
                float(st.session_state.ps_lat), 
                float(st.session_state.ps_lon), 
                distancia_teorica_1_fora_m / 1000.0, 
                az_para_1_fora
            )
            
            # FORÇA a atualização na memória do Streamlit para preencher as caixas automaticamente
            st.session_state["folder_lat_1_fora"] = formatar_lat_dm(lat_1_fora_calc)
            st.session_state["folder_lon_1_fora"] = formatar_lon_dm(lon_1_fora_calc)
        # -------------------------------------------------------------

        ref1, ref2, ref3 = st.columns(3)

        with ref1:
            st.markdown("#### 4' FORA")
            lat_4_fora = st.text_input("Latitude 4' fora", value="-", key="folder_lat_4_fora")
            lon_4_fora = st.text_input("Longitude 4' fora", value="-", key="folder_lon_4_fora")

        with ref2:
            st.markdown("#### 1' FORA")
            st.info(f"Distância teórica: {distancia_teorica_1_fora_m:,.0f} m".replace(",", "."))
            
            # Como forçamos a memória ali em cima, elas vão carregar preenchidas sozinhas!
            lat_1_fora = st.text_input("Latitude 1' fora", key="folder_lat_1_fora")
            lon_1_fora = st.text_input("Longitude 1' fora", key="folder_lon_1_fora")

        with ref3:
            st.markdown("#### PONTO DE SAÍDA - PS")
            if "ps_lat" in st.session_state and "ps_lon" in st.session_state:
                ps_lat_dm = formatar_lat_dm(float(st.session_state.ps_lat))
                ps_lon_dm = formatar_lon_dm(float(st.session_state.ps_lon))
                st.success("PS carregado automaticamente.")
                st.write(f"Latitude: **{ps_lat_dm}**")
                st.write(f"Longitude: **{ps_lon_dm}**")
            else:
                ps_lat_dm = ""
                ps_lon_dm = ""
                st.warning("PS ainda não registrado. Gere o KMZ ou registre o PS antes.")
        declinacao_ref = float(
            st.session_state.get(
                "declinacao_usada",
                st.session_state.get("declinacao", 0.0)
            ) or 0.0
        )

        direcao_base_verdadeira = float(
            st.session_state.get(
                "direcao_vento_verdadeira_kmz",
                0.0
            )
        )

        eixo_nariz = verdadeiro_para_magnetico(
            direcao_base_verdadeira,
            declinacao_ref
        )

        eixo_cauda = contra_azimute(eixo_nariz)

        if tipo_lancamento_folder == "Lançamento de Nariz":
            eixo_lancamento = eixo_nariz
        else:
            eixo_lancamento = eixo_cauda

        eixo_navegacao = eixo_cauda

        coord_zl = formatar_coord_dm(
            float(st.session_state.lat),
            float(st.session_state.lon)
        )

        altitude_zl_ft = float(st.session_state.get("altitude_ft", 0.0))
        altura_comandamento_ft = float(st.session_state.get("altura_comandamento_ft", 0.0))

        altitude_aerodromo_ft = st.session_state.get("altitude_aerodromo_partida_ft")
        altitude_ps_ft = st.session_state.get("altitude_consulta_ft")
        qfe_hpa = st.session_state.get("qfe_consulta_hpa")
        ajuste_altimetro_ft = st.session_state.get("altimetro_aerodromo_alvo_ft")

        def fmt_ft(valor):
            if valor is None:
                return ""
            return f"{valor:,.0f} ft".replace(",", "X").replace(".", ",").replace("X", ".")

        def fmt_qfe(valor):
            if valor is None:
                return ""
            return f"{valor:.0f}".zfill(4)

        dados_folder = {
            "localidade": st.session_state.get("localidade_alvo", ""),
            "coord_zl": coord_zl,
            "eixo_lancamento": f"{eixo_lancamento:.0f}°",
            "eixo_navegacao": f"{eixo_navegacao:.0f}°",
            "altura_comandamento_ft": fmt_ft(altura_comandamento_ft),
            "velocidade_anv": f"{velocidade_anv:.0f} m/s",
            "distancia_1_fora": f"{distancia_teorica_1_fora_m:,.0f} m".replace(",", "."),
            "altitude_zl_ft": fmt_ft(altitude_zl_ft),
            "altitude_aerodromo_ft": fmt_ft(altitude_aerodromo_ft),
            "altitude_ps_ft": fmt_ft(altitude_ps_ft),
            "daa_qfe": fmt_qfe(qfe_hpa),
            "ajuste_altimetro": (
                f"{ajuste_altimetro_ft:+.0f}".replace(".", ",")
                if ajuste_altimetro_ft is not None
                else ""
            ),
            "lat_4_fora": lat_4_fora,
            "lon_4_fora": lon_4_fora,
            "lat_1_fora": lat_1_fora,
            "lon_1_fora": lon_1_fora,
            "ps_lat_dm": ps_lat_dm,
            "ps_lon_dm": ps_lon_dm,
        }
        st.markdown("### Prévia dos dados")

        c1, c2 = st.columns(2)

        with c1:
            st.write(f"**Coordenada ZL:** {dados_folder['coord_zl']}")
            st.write(f"**Eixo de lançamento:** {dados_folder['eixo_lancamento']}")
            st.write(f"**Eixo de navegação:** {dados_folder['eixo_navegacao']}")
            st.write(f"**Alt comandamento:** {dados_folder['altura_comandamento_ft']}")
            st.write(f"**Velocidade da Anv:** {dados_folder['velocidade_anv']}")
            st.write(f"**Distância teórica 1' fora:** {dados_folder['distancia_1_fora']}")
        with c2:
            st.write(f"**Altitude da ZL:** {dados_folder['altitude_zl_ft']}")
            st.write(f"**Altitude Adrm:** {dados_folder['altitude_aerodromo_ft']}")
            st.write(
                f"**Altitude PS / DAA:** "
                f"{dados_folder['altitude_ps_ft']} / {dados_folder['daa_qfe']}"
            )
            st.write(f"**Ajuste de altímetro:** {dados_folder['ajuste_altimetro']}")
            st.write("**Pqdt embarcados:** ")

        gerar_acionado = st.button("📄 Gerar Folder do Piloto", type="primary")

        if gerar_acionado:
            registrar_log_missao("Gerou Folder do Piloto (Word)") # <--- RASTREADOR
            with st.spinner("📄 Gerando arquivo Word..."):                # Gera o arquivo em memória instantaneamente (sem satélite)
                st.session_state.docx_pronto = gerar_folder_piloto_docx(dados_folder)
            st.success("✅ Folder gerado com sucesso!")

        if "docx_pronto" in st.session_state:
            st.download_button(
                label="⬇️ Baixar Folder do Piloto em Word",
                data=st.session_state.docx_pronto,
                file_name="Folder_do_Piloto.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_download_folder_final"
            )
# =====================================================
# BARRA LATERAL: EXPORTAÇÃO DA MISSÃO (BACKUP)
# =====================================================

st.sidebar.markdown("---")
st.sidebar.markdown("### 💾 Arquivo de Missão")
st.sidebar.caption("Exporte todos os dados calculados para arquivamento ou análise pós-salto (AAR).")

def gerar_relatorio_txt():
    import datetime
    agora = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    # Busca os dados salvos na memória, se não existirem coloca 'N/A'
    lat = st.session_state.get('lat', 'N/A')
    lon = st.session_state.get('lon', 'N/A')
    alt = st.session_state.get('altitude_ft', 'N/A')
    dec = st.session_state.get('declinacao_usada', st.session_state.get('declinacao', 'N/A'))
    
    vento = st.session_state.get('vento_medio_velame', 'N/A')
    dir_pond = st.session_state.get('direcao_vento_verdadeira_kmz', 'N/A')
    
    dist_nm = st.session_state.get('distancia_velame_aberto_nm', 'N/A')
    ctehz = st.session_state.get('ctehz_velame_aberto_kt', 'N/A')
    k_val = st.session_state.get('k_velame_aberto_kft_h', 'N/A')
    fs_val = st.session_state.get('fs_velame_aberto_kft', 'N/A')

    linhas = [
        "=====================================================",
        "      RELATÓRIO TÁTICO DE PLANEJAMENTO - SLOP        ",
        "=====================================================",
        f"Data da exportação: {agora}",
        "",
        "[ 1. DADOS DO ALVO ]",
        f"Localidade / ZL: {st.session_state.get('localidade_alvo', 'N/A')}",
        f"Latitude:  {lat}",
        f"Longitude: {lon}",
        f"Altitude ZL (ft): {alt}",
        f"Declinação Magnética Aplicada: {dec}°",
        "",
        "[ 2. DADOS DO VENTO (VELAME ABERTO) ]",
        f"Vento Médio: {vento} kt",
        f"Direção Ponderada: {dir_pond}°",
        "",
        "[ 3. PARÂMETROS DE NAVEGAÇÃO GERAL ]",
        f"Distância D (NM): {dist_nm}",
        f"Constante Horizontal (CteHz): {ctehz} kt",
        f"Constante Vertical (K): {k_val} kft/h",
        f"Fator de Segurança (FS): {fs_val} kft",
        "",
        "[ 4. PONTOS DE CONTROLE REGISTRADOS ]"
    ]

    # Adicionar a tabela de pontos de controle se houver
    pontos = st.session_state.get("pontos_controle", [])
    if not pontos:
        linhas.append("Nenhum ponto registrado.")
    else:
        for p in pontos:
            obs = p.get('Observação', '')
            linhas.append(f"- Ponto {p['ID']} | Dist: {p['Dist (km)']} km | Alt: {p['Alt P Ctle (ft)']} ft | Vento: {p['Vento Médio (kt)']} kt | Obs: {obs}")

    # =====================================================
    # NOVO BLOCO: PLANEJAMENTO DKVA
    # =====================================================
    linhas.append("")
    linhas.append("[ 5. PLANEJAMENTO DKVA (SALTO SOBRE O ALVO) ]")
    
    dkva_a = st.session_state.get('dkva_a_kft', 'N/A')
    tipo_lanc_dkva = st.session_state.get('tipo_lanc_dkva_radio', 'N/A')
    
    if dkva_a != 'N/A' and isinstance(vento, (int, float)):
        d_dkva_m = 25 * vento * dkva_a
        eixo_nav = (dir_pond + 180) % 360 if isinstance(dir_pond, (int, float)) else 'N/A'
        
        linhas.append(f"Constante K: 25")
        linhas.append(f"Altura de Abertura (A): {dkva_a} kft")
        linhas.append(f"Distância Calculada (D): {d_dkva_m:.0f} metros")
        linhas.append(f"Eixo de Navegação (Nariz): {eixo_nav:.0f}°" if isinstance(eixo_nav, (int, float)) else f"Eixo de Navegação: {eixo_nav}")
        linhas.append(f"Tipo de Lançamento: {tipo_lanc_dkva}")
        
        if tipo_lanc_dkva in ["Lançamento de Cauda", "Lançamento Boca do Cone"]:
            nb = st.session_state.get('nb_dkva', 0)
            ib = st.session_state.get('ib_dkva', 0.0)
            ae = st.session_state.get('ae_dkva', 'N/A')
            linhas.append(f"Aeronave: {ae} | Blocos: {nb} | Intervalo: {ib} s")
    else:
        linhas.append("DKVA não preenchido nesta sessão.")

    # =====================================================
    
    linhas.append("")
    linhas.append("[ 6. WINDGRAM TEXTUAL UTILIZADO (CÓPIA BRUTA) ]")
    linhas.append(st.session_state.get("windgram_texto", "Nenhum windgram colado."))
    linhas.append("=====================================================")

    return "\n".join(linhas)

# Gera o texto e cria o botão de download
relatorio_str = gerar_relatorio_txt()

st.sidebar.download_button(
    label="📥 Baixar Relatório (TXT)",
    data=relatorio_str,
    file_name="Relatorio_Missao_SLOP.txt",
    mime="text/plain",
    use_container_width=True
)